"""
Generate a comprehensive Word document containing all project contents
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code, language=""):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    # Light gray background effect
    return para

def read_file_safe(filepath):
    """Safely read file contents"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        try:
            with open(filepath, 'r', encoding='latin-1') as f:
                return f.read()
        except:
            return "[Binary file or unable to read]"

def create_project_documentation():
    """Create comprehensive project documentation"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========== TITLE PAGE ==========
    title = doc.add_heading('Stealth Keylogger Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 0, 128)
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    add_heading_with_style(doc, 'Table of Contents', 1)
    
    toc_items = [
        "1. Project Overview",
        "2. Features and Capabilities",
        "3. Technical Architecture",
        "4. Installation and Setup",
        "5. Source Code Documentation",
        "   5.1 Main Application (main.py)",
        "   5.2 Keylogger Module (logger.py)",
        "   5.3 Input Processor (processor.py)",
        "   5.4 Encryption Utilities (crypto_utils.py)",
        "   5.5 Network Sender (network_sender.py)",
        "   5.6 Face Authentication (face_auth.py)",
        "   5.7 System Information Module (system_info.py)",
        "   5.8 Log Viewer (view_logs.py)",
        "   5.9 Security Setup (setup_security.py)",
        "6. Build Instructions",
        "7. System Information Module",
        "8. Requirements and Dependencies",
        "9. Ethical Considerations",
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 1. PROJECT OVERVIEW ==========
    add_heading_with_style(doc, '1. Project Overview', 1)
    
    doc.add_paragraph(
        'This project is a comprehensive educational keylogger disguised as a Minecraft Launcher. '
        'It demonstrates advanced cybersecurity concepts including keystroke capture, encryption, '
        'remote logging, and biometric authentication.'
    )
    
    doc.add_paragraph(
        'The project is designed for educational purposes to understand how malware operates, '
        'implement security measures, and learn about system reconnaissance techniques.'
    )
    
    # ========== 2. FEATURES ==========
    add_heading_with_style(doc, '2. Features and Capabilities', 1)
    
    features = [
        ('Stealth Installation', 'Disguised as Minecraft_Setup.bat and Minecraft_Launcher.exe'),
        ('Remote Logging', 'Sends logs to Discord webhook every 30 seconds (configurable)'),
        ('Secure Storage', 'AES-256 encryption using Fernet for all captured data'),
        ('Face Authentication', 'Biometric security using OpenCV and LBPH face recognition'),
        ('Persistence', 'Automatically runs on system startup via Registry (Windows)'),
        ('System Reconnaissance', 'Comprehensive system information gathering module'),
        ('Multi-Platform Support', 'Windows, Linux, and Android (via Buildozer)'),
        ('Cross-Platform Build', 'Separate build scripts for each platform'),
    ]
    
    for feature, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature}: ').bold = True
        p.add_run(description)
    
    # ========== 3. TECHNICAL ARCHITECTURE ==========
    add_heading_with_style(doc, '3. Technical Architecture', 1)
    
    add_heading_with_style(doc, '3.1 Technology Stack', 2)
    
    tech_stack = [
        ('Language', 'Python 3.13'),
        ('Key Capture', 'pynput library'),
        ('Encryption', 'cryptography (Fernet/AES-256)'),
        ('Face Recognition', 'opencv-python with Haar Cascades and LBPH'),
        ('Network Communication', 'urllib (built-in, no external dependencies)'),
        ('Build Tool', 'PyInstaller for executable creation'),
    ]
    
    for tech, desc in tech_stack:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{tech}: ').bold = True
        p.add_run(desc)
    
    add_heading_with_style(doc, '3.2 Module Architecture', 2)
    
    modules = [
        ('main.py', 'Main application entry point, orchestrates all components'),
        ('logger.py', 'Keyboard event capture using pynput'),
        ('processor.py', 'Processes and formats captured keystrokes'),
        ('crypto_utils.py', 'Encryption/decryption utilities using Fernet'),
        ('network_sender.py', 'Sends logs to Discord webhook'),
        ('face_auth.py', 'Face recognition for authentication'),
        ('system_info.py', 'System reconnaissance and information gathering'),
        ('view_logs.py', 'Secure log viewer with face authentication'),
        ('setup_security.py', 'Initial setup for encryption keys and face enrollment'),
    ]
    
    for module, desc in modules:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(module).bold = True
        p.add_run(f': {desc}')
    
    doc.add_page_break()
    
    # ========== 4. INSTALLATION AND SETUP ==========
    add_heading_with_style(doc, '4. Installation and Setup', 1)
    
    doc.add_paragraph('The project includes automated setup scripts for different platforms:')
    
    add_heading_with_style(doc, '4.1 Windows Installation', 2)
    doc.add_paragraph('1. Run build_portable.bat to create the executable')
    doc.add_paragraph('2. The release/ folder contains all deployment files')
    doc.add_paragraph('3. Run Minecraft_Setup.bat on target system')
    doc.add_paragraph('4. The keylogger installs to %APPDATA%\\Minecraft_Updater')
    
    add_heading_with_style(doc, '4.2 Linux Installation', 2)
    doc.add_paragraph('1. Run build_portable_linux.sh to create the binary')
    doc.add_paragraph('2. The release_linux/ folder contains deployment files')
    doc.add_paragraph('3. Run Minecraft_Setup.sh on target system')
    doc.add_paragraph('4. The keylogger installs to ~/.minecraft_updater')
    
    doc.add_page_break()
    
    # ========== 5. SOURCE CODE DOCUMENTATION ==========
    add_heading_with_style(doc, '5. Source Code Documentation', 1)
    
    # 5.1 Main Application
    add_heading_with_style(doc, '5.1 Main Application (main.py)', 2)
    doc.add_paragraph(
        'The main application file orchestrates all components of the keylogger. '
        'It initializes encryption, starts the key logger, and handles periodic log transmission.'
    )
    main_code = read_file_safe('main.py')
    add_code_block(doc, main_code, 'python')
    
    doc.add_page_break()
    
    # 5.2 Keylogger Module
    add_heading_with_style(doc, '5.2 Keylogger Module (logger.py)', 2)
    doc.add_paragraph(
        'The logger module uses pynput to capture keyboard events. '
        'It provides a clean interface for starting and stopping the key capture process.'
    )
    logger_code = read_file_safe('logger.py')
    add_code_block(doc, logger_code, 'python')
    
    doc.add_page_break()
    
    # 5.3 Input Processor
    add_heading_with_style(doc, '5.3 Input Processor (processor.py)', 2)
    doc.add_paragraph(
        'The processor module handles the conversion of raw keyboard events into readable text. '
        'It manages special keys like backspace, enter, and arrow keys.'
    )
    processor_code = read_file_safe('processor.py')
    add_code_block(doc, processor_code, 'python')
    
    doc.add_page_break()
    
    # 5.4 Encryption Utilities
    add_heading_with_style(doc, '5.4 Encryption Utilities (crypto_utils.py)', 2)
    doc.add_paragraph(
        'This module provides AES-256 encryption using the Fernet symmetric encryption scheme. '
        'All captured logs are encrypted before storage.'
    )
    crypto_code = read_file_safe('crypto_utils.py')
    add_code_block(doc, crypto_code, 'python')
    
    doc.add_page_break()
    
    # 5.5 Network Sender
    add_heading_with_style(doc, '5.5 Network Sender (network_sender.py)', 2)
    doc.add_paragraph(
        'The network sender module transmits captured logs to a Discord webhook. '
        'It uses only built-in urllib to avoid external dependencies.'
    )
    network_code = read_file_safe('network_sender.py')
    add_code_block(doc, network_code, 'python')
    
    doc.add_page_break()
    
    # 5.6 Face Authentication
    add_heading_with_style(doc, '5.6 Face Authentication (face_auth.py)', 2)
    doc.add_paragraph(
        'Implements biometric authentication using OpenCV. Uses Haar Cascades for face detection '
        'and LBPH (Local Binary Patterns Histograms) for face recognition.'
    )
    face_code = read_file_safe('face_auth.py')
    add_code_block(doc, face_code, 'python')
    
    doc.add_page_break()
    
    # 5.7 System Information Module
    add_heading_with_style(doc, '5.7 System Information Module (system_info.py)', 2)
    doc.add_paragraph(
        'Comprehensive reconnaissance module that gathers detailed system information including '
        'hardware specs, network configuration, installed software, security status, and more.'
    )
    sysinfo_code = read_file_safe('system_info.py')
    add_code_block(doc, sysinfo_code, 'python')
    
    doc.add_page_break()
    
    # 5.8 Log Viewer
    add_heading_with_style(doc, '5.8 Log Viewer (view_logs.py)', 2)
    doc.add_paragraph(
        'Secure log viewing application that requires face authentication before decrypting '
        'and displaying captured logs. Supports export to TXT, JSON, and PDF formats.'
    )
    view_code = read_file_safe('view_logs.py')
    add_code_block(doc, view_code, 'python')
    
    doc.add_page_break()
    
    # 5.9 Security Setup
    add_heading_with_style(doc, '5.9 Security Setup (setup_security.py)', 2)
    doc.add_paragraph(
        'Initial setup script that generates encryption keys and enrolls the user\'s face '
        'for authentication purposes.'
    )
    setup_code = read_file_safe('setup_security.py')
    add_code_block(doc, setup_code, 'python')
    
    doc.add_page_break()
    
    # ========== 6. BUILD INSTRUCTIONS ==========
    add_heading_with_style(doc, '6. Build Instructions', 1)
    
    build_instructions = read_file_safe('BUILD_INSTRUCTIONS.md')
    doc.add_paragraph(build_instructions)
    
    doc.add_page_break()
    
    # ========== 7. SYSTEM INFORMATION MODULE ==========
    add_heading_with_style(doc, '7. System Information Module', 1)
    
    sysinfo_readme = read_file_safe('SYSTEM_INFO_README.md')
    doc.add_paragraph(sysinfo_readme)
    
    doc.add_page_break()
    
    # ========== 8. REQUIREMENTS ==========
    add_heading_with_style(doc, '8. Requirements and Dependencies', 1)
    
    doc.add_paragraph('The project requires the following Python packages:')
    
    requirements = read_file_safe('requirements.txt')
    add_code_block(doc, requirements)
    
    doc.add_paragraph('\nInstallation command:')
    add_code_block(doc, 'pip install -r requirements.txt')
    
    # ========== 9. ETHICAL CONSIDERATIONS ==========
    add_heading_with_style(doc, '9. Ethical Considerations and Disclaimer', 1)
    
    doc.add_paragraph(
        'This project is developed strictly for EDUCATIONAL PURPOSES ONLY. '
        'It is designed to help cybersecurity students and professionals understand:'
    )
    
    educational_points = [
        'How keyloggers operate at a technical level',
        'Encryption and secure data storage techniques',
        'Biometric authentication implementation',
        'System reconnaissance methods used by malware',
        'Network communication and data exfiltration',
        'Multi-platform software development',
    ]
    
    for point in educational_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    warning = doc.add_paragraph()
    warning_run = warning.add_run('⚠️ WARNING: ')
    warning_run.bold = True
    warning_run.font.color.rgb = RGBColor(255, 0, 0)
    warning_run.font.size = Pt(14)
    
    warning.add_run(
        'Using this software on computers that you do not own or have explicit '
        'written permission to monitor is ILLEGAL and UNETHICAL. Unauthorized use '
        'may result in criminal prosecution under computer fraud and abuse laws.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The developers of this project assume NO LIABILITY for misuse of this software. '
        'By using this project, you agree to use it only in authorized, controlled '
        'environments for legitimate educational or research purposes.'
    )
    
    # ========== FOOTER ==========
    doc.add_page_break()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('End of Documentation')
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save the document
    output_file = 'Complete_Project_Documentation.docx'
    doc.save(output_file)
    print(f"✓ Documentation generated successfully: {output_file}")
    print(f"✓ Total pages: Multiple sections with complete source code")
    print(f"✓ File size: {os.path.getsize(output_file) / 1024:.2f} KB")
    
    return output_file

if __name__ == "__main__":
    print("Generating comprehensive project documentation...")
    print("=" * 60)
    create_project_documentation()
    print("=" * 60)
    print("Done!")
