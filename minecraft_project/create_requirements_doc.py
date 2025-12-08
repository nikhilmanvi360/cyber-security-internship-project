"""
Script to create Word document with Requirements (Hardware & Software)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Main Title
title = doc.add_heading('REQUIREMENTS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Hardware & Software', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a line break
doc.add_paragraph()

# ========== 1. DEVELOPMENT ENVIRONMENT REQUIREMENTS ==========
doc.add_heading('1. Development Environment Requirements', 1)

p1 = doc.add_paragraph(
    'The following hardware and software requirements are necessary for developing, building, '
    'and testing the multi-platform keylogger system.'
)

# 1.1 Hardware Requirements
doc.add_heading('1.1 Hardware Requirements', 2)

# Development Machine
doc.add_heading('Development Machine (Minimum Specifications)', 3)

dev_hardware = [
    ('Processor', 'Intel Core i5 (6th gen) or AMD Ryzen 5 equivalent', 'Dual-core, 2.5 GHz or higher'),
    ('RAM', '8 GB DDR4', '16 GB recommended for Android development'),
    ('Storage', '256 GB SSD', '50 GB free space minimum for build tools and dependencies'),
    ('Display', '1920x1080 resolution', 'For development and testing GUI components'),
    ('Network', 'Stable internet connection', 'Required for package downloads and webhook testing'),
]

table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Light Grid Accent 1'

hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Minimum Requirement'
hdr_cells[2].text = 'Notes'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for component, minimum, notes in dev_hardware:
    row_cells = table1.add_row().cells
    row_cells[0].text = component
    row_cells[1].text = minimum
    row_cells[2].text = notes

# Additional Hardware
doc.add_paragraph()
doc.add_heading('Additional Hardware Components', 3)

additional_hardware = [
    'Webcam (720p or higher) - Required for facial recognition training and authentication',
    'Microphone - Optional, not used in current implementation',
    'USB Drive (8 GB+) - For transferring executables between test systems',
    'Android Device (for Android build testing) - API level 21+ (Android 5.0 Lollipop or higher)',
    'Multiple test devices - For cross-platform compatibility testing'
]

for item in additional_hardware:
    doc.add_paragraph(item, style='List Bullet')

# 1.2 Software Requirements
doc.add_heading('1.2 Software Requirements', 2)

# Operating System
doc.add_heading('Operating System Requirements', 3)

os_requirements = [
    ('Windows Development', 'Windows 10/11 (64-bit)', 'For building Windows executables with PyInstaller'),
    ('Linux Development', 'Ubuntu 20.04+ / Debian 11+ / Fedora 35+', 'For building Linux binaries; requires X11/Xorg'),
    ('Android Development', 'Ubuntu 20.04+ (recommended)', 'Buildozer works best on Linux; WSL2 possible on Windows'),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'

hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Platform'
hdr_cells2[1].text = 'Requirement'
hdr_cells2[2].text = 'Purpose'

for cell in hdr_cells2:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for platform, requirement, purpose in os_requirements:
    row_cells = table2.add_row().cells
    row_cells[0].text = platform
    row_cells[1].text = requirement
    row_cells[2].text = purpose

# ========== 2. CORE SOFTWARE DEPENDENCIES ==========
doc.add_paragraph()
doc.add_heading('2. Core Software Dependencies', 1)

# 2.1 Python Environment
doc.add_heading('2.1 Python Environment', 2)

p2 = doc.add_paragraph()
p2.add_run('Python Version: ').bold = True
p2.add_run('Python 3.8 or higher (3.9 recommended)')

p3 = doc.add_paragraph()
p3.add_run('Installation: ').bold = True
p3.add_run('Download from https://www.python.org/ or use system package manager')

python_notes = [
    'Ensure Python is added to system PATH during installation',
    'pip (Python package installer) must be included',
    'Virtual environment support recommended (venv or virtualenv)',
    'Python 3.11+ may have compatibility issues with some libraries'
]

for note in python_notes:
    doc.add_paragraph(note, style='List Bullet')

# 2.2 Python Libraries
doc.add_heading('2.2 Python Libraries and Packages', 2)

p4 = doc.add_paragraph('The following Python packages are required (install via pip):')

# Core Libraries Table
libraries = [
    ('pynput', '1.7.6', 'Keyboard and mouse event monitoring', 'Core functionality'),
    ('cryptography', '41.0.0+', 'Fernet encryption (AES-256)', 'Data encryption'),
    ('opencv-python', '4.8.0+', 'Computer vision and face detection', 'Face authentication'),
    ('opencv-contrib-python', '4.8.0+', 'LBPH face recognizer', 'Face recognition'),
    ('numpy', '1.24.0+', 'Numerical operations for OpenCV', 'Image processing'),
    ('Pillow', '10.0.0+', 'Image handling (optional)', 'Image utilities'),
]

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Light Grid Accent 1'

hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Package'
hdr_cells3[1].text = 'Version'
hdr_cells3[2].text = 'Purpose'
hdr_cells3[3].text = 'Category'

for cell in hdr_cells3:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for package, version, purpose, category in libraries:
    row_cells = table3.add_row().cells
    row_cells[0].text = package
    row_cells[1].text = version
    row_cells[2].text = purpose
    row_cells[3].text = category

# Installation Command
doc.add_paragraph()
p5 = doc.add_paragraph()
p5.add_run('Installation Command:').bold = True

code_para = doc.add_paragraph('pip install pynput cryptography opencv-python opencv-contrib-python numpy')
code_para.style = 'Intense Quote'

# ========== 3. BUILD TOOLS AND COMPILERS ==========
doc.add_heading('3. Build Tools and Compilers', 1)

# 3.1 PyInstaller (Windows & Linux)
doc.add_heading('3.1 PyInstaller (Windows & Linux Builds)', 2)

p6 = doc.add_paragraph()
p6.add_run('Version: ').bold = True
p6.add_run('PyInstaller 5.13.0 or higher')

p7 = doc.add_paragraph()
p7.add_run('Purpose: ').bold = True
p7.add_run('Packages Python applications into standalone executables')

p8 = doc.add_paragraph()
p8.add_run('Installation: ').bold = True

code_para2 = doc.add_paragraph('pip install pyinstaller')
code_para2.style = 'Intense Quote'

pyinstaller_features = [
    'Creates single-file executables (--onefile flag)',
    'Bundles all dependencies and data files',
    'Supports hidden imports for platform-specific modules',
    'Generates .spec files for reproducible builds',
    'No console window option (--noconsole flag)'
]

for feature in pyinstaller_features:
    doc.add_paragraph(feature, style='List Bullet')

# 3.2 Buildozer (Android)
doc.add_heading('3.2 Buildozer (Android Builds)', 2)

p9 = doc.add_paragraph()
p9.add_run('Version: ').bold = True
p9.add_run('Buildozer 1.5.0 or higher')

p10 = doc.add_paragraph()
p10.add_run('Purpose: ').bold = True
p10.add_run('Compiles Python applications to Android APK packages')

p11 = doc.add_paragraph()
p11.add_run('Platform: ').bold = True
p11.add_run('Linux only (Ubuntu recommended) or WSL2 on Windows')

p12 = doc.add_paragraph()
p12.add_run('Installation: ').bold = True

code_para3 = doc.add_paragraph('pip install buildozer\nsudo apt-get install -y git zip unzip openjdk-17-jdk autoconf libtool')
code_para3.style = 'Intense Quote'

buildozer_requirements = [
    'Java Development Kit (JDK) 17',
    'Android SDK (automatically downloaded by Buildozer)',
    'Android NDK r25b (automatically downloaded)',
    'Cython for Python-to-C compilation',
    'Git for dependency management',
    'Minimum 10 GB free space for Android SDK/NDK'
]

doc.add_paragraph('Additional Requirements:')
for req in buildozer_requirements:
    doc.add_paragraph(req, style='List Bullet')

# ========== 4. PLATFORM-SPECIFIC REQUIREMENTS ==========
doc.add_heading('4. Platform-Specific Requirements', 1)

# 4.1 Windows-Specific
doc.add_heading('4.1 Windows Platform Requirements', 2)

windows_reqs = [
    ('Visual C++ Redistributable', 'Microsoft Visual C++ 2015-2022 Redistributable', 'Required for some Python packages'),
    ('Windows SDK', 'Windows 10 SDK (optional)', 'For advanced Win32 API development'),
    ('Registry Access', 'Administrator privileges (for persistence)', 'To modify HKCU\\Software\\...\\Run registry keys'),
    ('PowerShell', 'PowerShell 5.1 or higher', 'For running installation scripts (.bat files)'),
]

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'

hdr_cells4 = table4.rows[0].cells
hdr_cells4[0].text = 'Component'
hdr_cells4[1].text = 'Requirement'
hdr_cells4[2].text = 'Purpose'

for cell in hdr_cells4:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for component, requirement, purpose in windows_reqs:
    row_cells = table4.add_row().cells
    row_cells[0].text = component
    row_cells[1].text = requirement
    row_cells[2].text = purpose

# 4.2 Linux-Specific
doc.add_paragraph()
doc.add_heading('4.2 Linux Platform Requirements', 2)

linux_reqs = [
    ('X11/Xorg', 'X Window System', 'Required for keyboard event capture (pynput)'),
    ('systemd', 'systemd init system', 'For creating user services (persistence)'),
    ('Python Dev Headers', 'python3-dev package', 'For compiling C extensions'),
    ('Build Essentials', 'build-essential package', 'GCC compiler and make tools'),
    ('GTK/Qt Libraries', 'libgtk-3-0 or qt5-default', 'For GUI components (if needed)'),
]

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Light Grid Accent 1'

hdr_cells5 = table5.rows[0].cells
hdr_cells5[0].text = 'Component'
hdr_cells5[1].text = 'Requirement'
hdr_cells5[2].text = 'Purpose'

for cell in hdr_cells5:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for component, requirement, purpose in linux_reqs:
    row_cells = table5.add_row().cells
    row_cells[0].text = component
    row_cells[1].text = requirement
    row_cells[2].text = purpose

p13 = doc.add_paragraph()
p13.add_run('Installation Command (Ubuntu/Debian):').bold = True

code_para4 = doc.add_paragraph('sudo apt-get install python3-dev build-essential libgtk-3-0 xorg')
code_para4.style = 'Intense Quote'

# 4.3 Android-Specific
doc.add_paragraph()
doc.add_heading('4.3 Android Platform Requirements', 2)

android_reqs = [
    ('Android Device', 'Android 5.0+ (API 21+)', 'Physical device or emulator for testing'),
    ('Developer Mode', 'USB debugging enabled', 'For ADB installation and testing'),
    ('Accessibility Service', 'User must enable manually', 'Required for keystroke capture'),
    ('Permissions', 'INTERNET, CAMERA, BOOT_COMPLETED', 'Declared in AndroidManifest.xml'),
    ('Storage', 'Minimum 100 MB free space', 'For app installation and logs'),
]

table6 = doc.add_table(rows=1, cols=3)
table6.style = 'Light Grid Accent 1'

hdr_cells6 = table6.rows[0].cells
hdr_cells6[0].text = 'Component'
hdr_cells6[1].text = 'Requirement'
hdr_cells6[2].text = 'Purpose'

for cell in hdr_cells6:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for component, requirement, purpose in android_reqs:
    row_cells = table6.add_row().cells
    row_cells[0].text = component
    row_cells[1].text = requirement
    row_cells[2].text = purpose

# ========== 5. NETWORK REQUIREMENTS ==========
doc.add_heading('5. Network Requirements', 1)

# 5.1 Internet Connectivity
doc.add_heading('5.1 Internet Connectivity', 2)

network_reqs = [
    'Stable internet connection for package downloads and dependency installation',
    'HTTPS access to Discord API (discord.com/api/webhooks/*)',
    'Outbound connections on port 443 (HTTPS)',
    'DNS resolution capability (for webhook domain resolution)',
    'Minimum bandwidth: 1 Mbps (for data exfiltration testing)',
    'Firewall exceptions for Python and compiled executables'
]

for req in network_reqs:
    doc.add_paragraph(req, style='List Bullet')

# 5.2 Discord Webhook
doc.add_heading('5.2 Discord Webhook Configuration', 2)

p14 = doc.add_paragraph()
p14.add_run('Requirement: ').bold = True
p14.add_run('Active Discord webhook URL for data exfiltration testing')

webhook_setup = [
    'Create a Discord server (free account)',
    'Create a text channel for receiving logs',
    'Generate webhook URL from channel settings',
    'Configure webhook URL in network_sender.py',
    'Test connectivity before deployment'
]

doc.add_paragraph('Setup Steps:')
for i, step in enumerate(webhook_setup, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

# ========== 6. DEVELOPMENT TOOLS (OPTIONAL) ==========
doc.add_heading('6. Development Tools (Recommended)', 1)

# 6.1 IDE and Editors
doc.add_heading('6.1 Integrated Development Environments', 2)

ides = [
    ('Visual Studio Code', 'Lightweight, Python extensions, integrated terminal', 'Recommended'),
    ('PyCharm Community', 'Full-featured Python IDE, debugging tools', 'Alternative'),
    ('Sublime Text', 'Fast text editor with Python support', 'Lightweight option'),
    ('Vim/Neovim', 'Terminal-based editor for Linux development', 'Advanced users'),
]

table7 = doc.add_table(rows=1, cols=3)
table7.style = 'Light Grid Accent 1'

hdr_cells7 = table7.rows[0].cells
hdr_cells7[0].text = 'Tool'
hdr_cells7[1].text = 'Features'
hdr_cells7[2].text = 'Category'

for cell in hdr_cells7:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for tool, features, category in ides:
    row_cells = table7.add_row().cells
    row_cells[0].text = tool
    row_cells[1].text = features
    row_cells[2].text = category

# 6.2 Version Control
doc.add_paragraph()
doc.add_heading('6.2 Version Control System', 2)

p15 = doc.add_paragraph()
p15.add_run('Git: ').bold = True
p15.add_run('Version 2.30 or higher for source code management')

git_features = [
    'Track code changes and version history',
    'Branch management for feature development',
    'Collaboration and code review (if team-based)',
    'Integration with GitHub/GitLab for backup'
]

for feature in git_features:
    doc.add_paragraph(feature, style='List Bullet')

# 6.3 Testing Tools
doc.add_paragraph()
doc.add_heading('6.3 Testing and Debugging Tools', 2)

testing_tools = [
    ('Process Explorer (Windows)', 'Monitor running processes and resource usage'),
    ('htop (Linux)', 'System monitoring and process management'),
    ('Wireshark', 'Network traffic analysis and packet inspection'),
    ('ADB (Android Debug Bridge)', 'Android device debugging and log viewing'),
    ('Python Debugger (pdb)', 'Built-in Python debugging tool'),
]

for tool, purpose in testing_tools:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{tool}: ').bold = True
    p.add_run(purpose)

# ========== 7. TARGET DEPLOYMENT REQUIREMENTS ==========
doc.add_heading('7. Target Deployment Requirements', 1)

p16 = doc.add_paragraph(
    'The following are the minimum requirements for systems where the keylogger will be deployed and tested:'
)

# 7.1 Windows Target
doc.add_heading('7.1 Windows Target Systems', 2)

windows_target = [
    'Operating System: Windows 10 (1809+) or Windows 11',
    'Architecture: 64-bit (x64)',
    'RAM: 4 GB minimum',
    'Storage: 100 MB free space',
    'Permissions: Standard user account (no admin required)',
    'Network: Internet connection for data exfiltration'
]

for req in windows_target:
    doc.add_paragraph(req, style='List Bullet')

# 7.2 Linux Target
doc.add_heading('7.2 Linux Target Systems', 2)

linux_target = [
    'Operating System: Ubuntu 18.04+, Debian 10+, Fedora 32+, or compatible',
    'Desktop Environment: X11/Xorg (Wayland not supported by pynput)',
    'Architecture: 64-bit (x86_64)',
    'RAM: 2 GB minimum',
    'Storage: 50 MB free space',
    'Permissions: Standard user account',
    'Network: Internet connection for data exfiltration'
]

for req in linux_target:
    doc.add_paragraph(req, style='List Bullet')

# 7.3 Android Target
doc.add_heading('7.3 Android Target Devices', 2)

android_target = [
    'Operating System: Android 5.0 (Lollipop) or higher',
    'API Level: 21+',
    'Architecture: ARM or ARM64',
    'RAM: 2 GB minimum',
    'Storage: 100 MB free space',
    'Permissions: User must grant Accessibility Service access',
    'Network: Wi-Fi or mobile data connection'
]

for req in android_target:
    doc.add_paragraph(req, style='List Bullet')

# ========== 8. SECURITY REQUIREMENTS ==========
doc.add_heading('8. Security and Testing Requirements', 1)

# 8.1 Isolated Testing Environment
doc.add_heading('8.1 Isolated Testing Environment', 2)

p17 = doc.add_paragraph()
p17.add_run('CRITICAL: ').bold = True
p17.add_run('All testing must be conducted in isolated, controlled environments.')

isolation_requirements = [
    'Dedicated test machines (not primary work computers)',
    'Virtual machines (VirtualBox, VMware) for safe testing',
    'No sensitive data on test systems',
    'Network isolation or monitoring for exfiltration testing',
    'Ability to restore systems to clean state (snapshots)',
    'No connection to production networks'
]

for req in isolation_requirements:
    doc.add_paragraph(req, style='List Bullet')

# 8.2 Legal and Ethical Requirements
doc.add_heading('8.2 Legal and Ethical Requirements', 2)

legal_requirements = [
    'Written authorization for all test systems',
    'Ownership or explicit permission for all devices',
    'Compliance with local computer crime laws',
    'No deployment on unauthorized systems',
    'Secure storage of source code and executables',
    'Documentation of educational purpose'
]

for req in legal_requirements:
    doc.add_paragraph(req, style='List Bullet')

# ========== 9. SUMMARY TABLE ==========
doc.add_heading('9. Quick Reference Summary', 1)

p18 = doc.add_paragraph('Complete requirements summary for quick reference:')

# Summary Table
summary_data = [
    ('Development OS', 'Windows 10/11 or Ubuntu 20.04+', 'For building executables'),
    ('Python Version', '3.8 - 3.10 (3.9 recommended)', 'Core programming language'),
    ('RAM', '8 GB minimum, 16 GB recommended', 'Development and compilation'),
    ('Storage', '50 GB free space', 'SDK, NDK, dependencies'),
    ('Webcam', '720p or higher', 'Face authentication'),
    ('Internet', 'Stable broadband connection', 'Package downloads, testing'),
    ('Key Libraries', 'pynput, cryptography, opencv-python', 'Core functionality'),
    ('Build Tools', 'PyInstaller 5.13+, Buildozer 1.5+', 'Executable compilation'),
    ('Android Device', 'API 21+ (Android 5.0+)', 'Android testing'),
    ('Test Environment', 'Isolated VMs or dedicated hardware', 'Safe testing'),
]

table8 = doc.add_table(rows=1, cols=3)
table8.style = 'Medium Grid 1 Accent 1'

hdr_cells8 = table8.rows[0].cells
hdr_cells8[0].text = 'Category'
hdr_cells8[1].text = 'Requirement'
hdr_cells8[2].text = 'Purpose'

for cell in hdr_cells8:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for category, requirement, purpose in summary_data:
    row_cells = table8.add_row().cells
    row_cells[0].text = category
    row_cells[1].text = requirement
    row_cells[2].text = purpose

# ========== CONCLUSION ==========
doc.add_paragraph()
doc.add_heading('Conclusion', 1)

p19 = doc.add_paragraph(
    'This comprehensive requirements specification ensures that all necessary hardware and software '
    'components are in place for successful development, compilation, and testing of the multi-platform '
    'keylogger system. Meeting these requirements will enable smooth development workflow and reliable '
    'cross-platform deployment.'
)

p20 = doc.add_paragraph()
p20.add_run('Note: ').bold = True
p20.add_run(
    'All requirements are subject to change based on library updates and platform evolution. '
    'Always refer to the latest documentation for each dependency and tool.'
)

# Save the document
doc.save('Requirements_Hardware_and_Software.docx')
print("✓ Word document created successfully: Requirements_Hardware_and_Software.docx")
print("\nDocument includes:")
print("  • Development Environment Requirements (Hardware & Software)")
print("  • Core Software Dependencies (Python libraries)")
print("  • Build Tools and Compilers (PyInstaller, Buildozer)")
print("  • Platform-Specific Requirements (Windows, Linux, Android)")
print("  • Network Requirements (Discord webhook)")
print("  • Development Tools (IDEs, version control, testing)")
print("  • Target Deployment Requirements")
print("  • Security and Testing Requirements")
print("  • Quick Reference Summary Table")
