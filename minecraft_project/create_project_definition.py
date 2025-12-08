"""
Script to create Word document with Project Definition (Purpose and Scope)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Main Title
title = doc.add_heading('PROJECT DEFINITION', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Purpose and Scope', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a line break
doc.add_paragraph()

# ========== 1. PROJECT PURPOSE ==========
doc.add_heading('1. Project Purpose', 1)

# 1.1 Primary Purpose
doc.add_heading('1.1 Primary Purpose', 2)
p1 = doc.add_paragraph(
    'The primary purpose of this cybersecurity internship project is to develop a comprehensive understanding '
    'of offensive security techniques through the practical implementation of an advanced keylogging system. '
    'This project serves as an educational tool to demonstrate:'
)

primary_purposes = [
    'How malicious actors design and deploy keystroke logging malware',
    'The technical mechanisms behind data exfiltration and covert communication',
    'Encryption techniques used to protect stolen data from forensic analysis',
    'Persistence mechanisms that allow malware to survive system reboots and security scans',
    'Cross-platform malware development strategies for Windows, Linux, and Android',
    'Social engineering tactics used to disguise malicious software as legitimate applications'
]

for purpose in primary_purposes:
    doc.add_paragraph(purpose, style='List Bullet')

# 1.2 Educational Objectives
doc.add_heading('1.2 Educational Objectives', 2)
p2 = doc.add_paragraph('This project aims to provide hands-on learning in the following areas:')

educational_objectives = [
    ('Offensive Security Research', 'Understanding attack vectors, exploitation techniques, and malware behavior patterns'),
    ('Defensive Cybersecurity', 'Learning to identify, detect, and mitigate keylogging threats through reverse engineering'),
    ('Cryptographic Implementation', 'Practical application of AES-256 encryption, key management, and secure data handling'),
    ('System-Level Programming', 'Working with low-level APIs, keyboard hooks, and operating system internals'),
    ('Ethical Hacking Principles', 'Applying responsible disclosure practices and understanding legal boundaries'),
    ('Penetration Testing Skills', 'Simulating real-world attack scenarios in controlled environments')
]

for obj_name, description in educational_objectives:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{obj_name}: ').bold = True
    p.add_run(description)

# 1.3 Research Questions
doc.add_heading('1.3 Research Questions', 2)
p3 = doc.add_paragraph('This project seeks to answer the following research questions:')

research_questions = [
    'How can keystroke logging malware evade modern antivirus and endpoint detection systems?',
    'What are the most effective persistence mechanisms across different operating systems?',
    'How do attackers implement covert data exfiltration using legitimate communication channels?',
    'What role does encryption play in protecting malware operations from security analysis?',
    'How can biometric authentication be integrated into malware for access control?',
    'What are the technical challenges in developing cross-platform malicious software?',
    'What defensive countermeasures are most effective against keylogging attacks?'
]

for i, question in enumerate(research_questions, 1):
    doc.add_paragraph(f'{i}. {question}', style='List Number')

# ========== 2. PROJECT SCOPE ==========
doc.add_heading('2. Project Scope', 1)

# 2.1 In-Scope Components
doc.add_heading('2.1 In-Scope Components', 2)
p4 = doc.add_paragraph('The following components and features are included within the project scope:')

# Core Functionality
doc.add_heading('Core Functionality', 3)
core_features = [
    'Real-time keystroke capture using cross-platform keyboard hooks',
    'Intelligent input processing with special key handling (Backspace, Enter, Arrow keys)',
    'AES-256 encryption for secure log storage',
    'Automated data exfiltration via Discord webhooks',
    'Biometric authentication using facial recognition (OpenCV LBPH)',
    'Process management and lifecycle control'
]

for feature in core_features:
    doc.add_paragraph(feature, style='List Bullet')

# Platform Support
doc.add_heading('Platform Support', 3)
platform_features = [
    ('Windows', 'Standalone executable (.exe) with registry-based persistence'),
    ('Linux', 'ELF binary with systemd service integration'),
    ('Android', 'APK package with accessibility service and boot receiver')
]

for platform, description in platform_features:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{platform}: ').bold = True
    p.add_run(description)

# Security Features
doc.add_heading('Security Features', 3)
security_features = [
    'Fernet encryption (AES-256-CBC with HMAC-SHA256)',
    'Symmetric key generation and management',
    'Encrypted log file storage (captured_text.enc)',
    'Secure decryption interface for authorized access',
    'Face recognition-based authentication system'
]

for feature in security_features:
    doc.add_paragraph(feature, style='List Bullet')

# Stealth Mechanisms
doc.add_heading('Stealth Mechanisms', 3)
stealth_features = [
    'Hidden installation directories (AppData, dot-prefixed folders)',
    'Process disguise as "Minecraft Launcher"',
    'No visible GUI or console windows',
    'Minimal resource consumption (CPU, memory)',
    'HTTPS-encrypted network traffic to legitimate services'
]

for feature in stealth_features:
    doc.add_paragraph(feature, style='List Bullet')

# Deployment Capabilities
doc.add_heading('Deployment Capabilities', 3)
deployment_features = [
    'Automated installation scripts for all platforms',
    'Persistence configuration (autostart on boot)',
    'Portable executable packaging (PyInstaller, Buildozer)',
    'Transfer packages for cross-platform deployment',
    'Uninstallation and cleanup scripts'
]

for feature in deployment_features:
    doc.add_paragraph(feature, style='List Bullet')

# 2.2 Out-of-Scope Components
doc.add_heading('2.2 Out-of-Scope Components', 2)
p5 = doc.add_paragraph('The following features and capabilities are explicitly excluded from this project:')

out_of_scope = [
    ('Rootkit Functionality', 'No kernel-level hooks or driver-based stealth mechanisms'),
    ('Advanced Evasion', 'No anti-debugging, anti-VM, or sandbox detection techniques'),
    ('Payload Delivery', 'No exploit development or vulnerability exploitation'),
    ('Command & Control', 'No bidirectional C2 server or remote command execution'),
    ('Screen Capture', 'No screenshot or screen recording capabilities'),
    ('Webcam Access', 'No unauthorized camera activation (except for face auth)'),
    ('Network Sniffing', 'No packet capture or network traffic interception'),
    ('Privilege Escalation', 'No attempts to gain administrator/root access'),
    ('Self-Propagation', 'No worm-like spreading or network propagation'),
    ('Data Destruction', 'No file deletion, encryption (ransomware), or system damage')
]

for item_name, description in out_of_scope:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{item_name}: ').bold = True
    p.add_run(description)

# 2.3 Technical Boundaries
doc.add_heading('2.3 Technical Boundaries', 2)

# Programming Languages
doc.add_heading('Programming Languages', 3)
p6 = doc.add_paragraph()
p6.add_run('Primary Language: ').bold = True
p6.add_run('Python 3.x (for cross-platform compatibility and rapid development)')

# Dependencies and Libraries
doc.add_heading('Key Dependencies', 3)
dependencies = [
    ('pynput', 'Cross-platform keyboard and mouse event handling'),
    ('cryptography', 'Fernet encryption implementation (AES-256)'),
    ('opencv-python', 'Computer vision and facial recognition'),
    ('PyInstaller', 'Executable packaging for Windows and Linux'),
    ('Buildozer', 'Android APK compilation and packaging')
]

for lib, purpose in dependencies:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{lib}: ').bold = True
    p.add_run(purpose)

# Target Environments
doc.add_heading('Target Environments', 3)
environments = [
    'Windows 10/11 (64-bit)',
    'Linux distributions with X11/Xorg (Ubuntu, Debian, Fedora)',
    'Android 5.0+ (API level 21+) with accessibility services'
]

for env in environments:
    doc.add_paragraph(env, style='List Bullet')

# 2.4 Ethical and Legal Boundaries
doc.add_heading('2.4 Ethical and Legal Boundaries', 2)

p7 = doc.add_paragraph()
p7.add_run('IMPORTANT: ').bold = True
p7.add_run('This project is developed strictly for educational and research purposes within a controlled environment. '
           'The following ethical and legal boundaries are strictly observed:')

ethical_boundaries = [
    'Testing conducted ONLY on personally-owned devices with explicit consent',
    'No deployment on systems owned by others without written authorization',
    'No use for malicious purposes, unauthorized surveillance, or illegal activities',
    'Compliance with local, state, and federal computer crime laws',
    'Adherence to responsible disclosure practices for any discovered vulnerabilities',
    'Clear documentation of educational intent and cybersecurity research purpose',
    'No distribution to unauthorized parties or public repositories without disclaimers',
    'Immediate cessation of any activity upon legal or ethical concerns'
]

for boundary in ethical_boundaries:
    doc.add_paragraph(boundary, style='List Bullet')

# ========== 3. PROJECT DELIVERABLES ==========
doc.add_heading('3. Project Deliverables', 1)

deliverables = [
    ('Source Code', 'Complete Python source code for all modules (logger, processor, crypto, face auth, network sender)'),
    ('Compiled Executables', 'Platform-specific binaries: Windows .exe, Linux ELF, Android .apk'),
    ('Installation Scripts', 'Automated deployment scripts for all supported platforms'),
    ('Build Specifications', 'PyInstaller .spec files and Buildozer configuration'),
    ('Technical Documentation', 'Comprehensive report detailing architecture, implementation, and security analysis'),
    ('User Guides', 'README files with build instructions, installation procedures, and usage guidelines'),
    ('Security Analysis', 'Defensive countermeasures, detection methods, and removal instructions'),
    ('Presentation Materials', 'PowerPoint presentation summarizing key findings and demonstrations')
]

for deliverable, description in deliverables:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{deliverable}: ').bold = True
    p.add_run(description)

# ========== 4. SUCCESS CRITERIA ==========
doc.add_heading('4. Success Criteria', 1)

p8 = doc.add_paragraph('The project will be considered successful upon meeting the following criteria:')

# Technical Success Criteria
doc.add_heading('4.1 Technical Success Criteria', 2)
technical_criteria = [
    'Successful keystroke capture on all three target platforms (Windows, Linux, Android)',
    'Functional AES-256 encryption and decryption of captured logs',
    'Reliable data exfiltration via Discord webhooks with 95%+ success rate',
    'Persistence mechanisms functioning correctly across system reboots',
    'Face authentication system achieving 80%+ recognition accuracy',
    'Executable size under 50MB for efficient deployment',
    'CPU usage under 5% during normal operation',
    'Memory footprint under 100MB'
]

for i, criterion in enumerate(technical_criteria, 1):
    doc.add_paragraph(f'{i}. {criterion}', style='List Number')

# Educational Success Criteria
doc.add_heading('4.2 Educational Success Criteria', 2)
educational_criteria = [
    'Demonstrated understanding of low-level system APIs and keyboard hooks',
    'Practical implementation of cryptographic algorithms and secure key management',
    'Successful cross-platform software development and deployment',
    'Comprehensive documentation of attack vectors and defensive countermeasures',
    'Ability to analyze and reverse engineer malware behavior',
    'Understanding of ethical hacking principles and legal compliance'
]

for i, criterion in enumerate(educational_criteria, 1):
    doc.add_paragraph(f'{i}. {criterion}', style='List Number')

# ========== 5. PROJECT CONSTRAINTS ==========
doc.add_heading('5. Project Constraints', 1)

# Time Constraints
doc.add_heading('5.1 Time Constraints', 2)
p9 = doc.add_paragraph(
    'The project is designed to be completed within a typical internship timeframe, with the following phases:'
)

time_phases = [
    ('Research & Planning', '1-2 weeks - Literature review, architecture design, tool selection'),
    ('Core Development', '2-3 weeks - Implementation of keylogger, encryption, and exfiltration'),
    ('Platform Adaptation', '1-2 weeks - Cross-platform compilation and testing'),
    ('Security Analysis', '1 week - Detection methods, countermeasures, and ethical review'),
    ('Documentation', '1 week - Report writing, presentation creation, and code documentation')
]

for phase, description in time_phases:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{phase}: ').bold = True
    p.add_run(description)

# Resource Constraints
doc.add_heading('5.2 Resource Constraints', 2)
resource_constraints = [
    'Development limited to open-source libraries and free tools',
    'Testing conducted on personal hardware (no cloud infrastructure)',
    'Network exfiltration using free Discord webhook service',
    'No budget for commercial software licenses or paid services',
    'Limited to Python ecosystem for cross-platform compatibility'
]

for constraint in resource_constraints:
    doc.add_paragraph(constraint, style='List Bullet')

# Technical Constraints
doc.add_heading('5.3 Technical Constraints', 2)
technical_constraints = [
    'Python performance limitations compared to compiled languages (C/C++)',
    'PyInstaller executable size overhead (larger than native binaries)',
    'Android accessibility service requires manual user enablement',
    'No kernel-level access (user-space only)',
    'Discord webhook rate limiting (30 requests per minute)',
    'Face recognition accuracy dependent on lighting conditions',
    'Cross-platform compatibility limited by pynput library capabilities'
]

for constraint in technical_constraints:
    doc.add_paragraph(constraint, style='List Bullet')

# ========== 6. RISK ASSESSMENT ==========
doc.add_heading('6. Risk Assessment', 1)

p10 = doc.add_paragraph(
    'The following risks have been identified and mitigation strategies implemented:'
)

risks = [
    ('Legal Risk', 'Unauthorized use or distribution', 'Clear documentation, ethical guidelines, controlled testing environment'),
    ('Detection Risk', 'Antivirus flagging as malware', 'Expected behavior; educational disclaimer included in documentation'),
    ('Technical Risk', 'Platform compatibility issues', 'Extensive testing on multiple OS versions and configurations'),
    ('Ethical Risk', 'Misuse by unauthorized parties', 'No public distribution; secure storage of source code and executables'),
    ('Privacy Risk', 'Unintended data capture', 'Testing only on isolated systems with no sensitive information'),
    ('Security Risk', 'Encryption key compromise', 'Secure key storage; documentation of key management best practices')
]

# Create table for risks
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Risk Category'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Mitigation Strategy'

# Make header bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for risk_cat, desc, mitigation in risks:
    row_cells = table.add_row().cells
    row_cells[0].text = risk_cat
    row_cells[1].text = desc
    row_cells[2].text = mitigation

# ========== CONCLUSION ==========
doc.add_paragraph()  # Add spacing
doc.add_heading('Conclusion', 1)

p11 = doc.add_paragraph(
    'This project definition establishes a clear purpose and scope for the development of an educational '
    'keylogging system designed to enhance understanding of offensive security techniques, malware behavior, '
    'and defensive countermeasures. By maintaining strict ethical boundaries and focusing on controlled '
    'research environments, this project provides valuable hands-on experience in cybersecurity while '
    'adhering to legal and professional standards.'
)

p12 = doc.add_paragraph(
    'The comprehensive scope encompasses cross-platform development, cryptographic implementation, stealth '
    'mechanisms, and data exfiltration techniques, providing a holistic view of modern malware capabilities. '
    'Through this project, critical insights into both offensive and defensive cybersecurity practices are '
    'gained, contributing to the broader goal of improving digital security awareness and preparedness.'
)

# Save the document
doc.save('Project_Definition_Purpose_and_Scope.docx')
print("✓ Word document created successfully: Project_Definition_Purpose_and_Scope.docx")
print("\nDocument includes:")
print("  • Project Purpose (Primary Purpose, Educational Objectives, Research Questions)")
print("  • Project Scope (In-Scope, Out-of-Scope, Technical Boundaries, Ethical Boundaries)")
print("  • Project Deliverables")
print("  • Success Criteria (Technical & Educational)")
print("  • Project Constraints (Time, Resource, Technical)")
print("  • Risk Assessment (with mitigation strategies)")
