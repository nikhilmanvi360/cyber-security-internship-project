"""
Script to create Word document with project introduction
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Create a new Document
    doc = Document()
    
    # Title
    title = doc.add_heading('Introduction', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Overview
    doc.add_heading('Project Overview', 1)
    p1 = doc.add_paragraph(
        'This cybersecurity internship project presents the development and analysis of an advanced, '
        'cross-platform keylogger system designed for educational purposes and penetration testing research. '
        'The application, strategically disguised as a "Minecraft Launcher," demonstrates sophisticated '
        'techniques employed by modern malware while providing insights into defensive countermeasures.'
    )
    
    # Project Objectives
    doc.add_heading('Project Objectives', 1)
    p2 = doc.add_paragraph('The primary objectives of this project were to:')
    
    objectives = [
        'Understand offensive security techniques by implementing real-world malware capabilities in a controlled environment',
        'Develop cross-platform expertise by creating deployable versions for Windows, Linux, and Android operating systems',
        'Implement military-grade encryption (AES-256) to secure captured data and demonstrate cryptographic principles',
        'Explore stealth and persistence mechanisms used by advanced persistent threats (APTs)',
        'Analyze detection methods and develop defensive countermeasures against keylogging attacks',
        'Apply ethical hacking principles within a legal and educational framework'
    ]
    
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(obj, style='List Number')
    
    # Technical Scope
    doc.add_heading('Technical Scope', 1)
    p3 = doc.add_paragraph(
        'The project encompasses approximately 563 lines of core Python code organized into seven modular components:'
    )
    
    components = [
        ('KeyLogger Engine', 'logger.py', 'Cross-platform keystroke capture using event-driven architecture'),
        ('Input Processor', 'processor.py', 'Intelligent keystroke parsing and buffer management'),
        ('Encryption Module', 'crypto_utils.py', 'Fernet-based AES-256 encryption implementation'),
        ('Face Authentication', 'face_auth.py', 'Biometric access control using OpenCV and LBPH recognition'),
        ('Network Exfiltration', 'network_sender.py', 'Covert data transmission via Discord webhooks'),
        ('Main Controller', 'main.py', 'Orchestration and lifecycle management'),
        ('Log Viewer', 'view_logs.py', 'Secure decryption and log analysis interface')
    ]
    
    for comp_name, file_name, description in components:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{comp_name}').bold = True
        p.add_run(f' ({file_name}) - {description}')
    
    # Key Features
    doc.add_heading('Key Features', 1)
    p4 = doc.add_paragraph('The system demonstrates several advanced capabilities:')
    
    features = [
        ('Multi-Platform Support', 'Deployable executables for Windows (.exe), Linux (ELF binary), and Android (.apk)'),
        ('Military-Grade Encryption', 'AES-256-CBC with HMAC-SHA256 authentication'),
        ('Biometric Security', 'Facial recognition-based access control using Local Binary Patterns Histograms (LBPH)'),
        ('Stealth Operations', 'Hidden installation directories, registry persistence, and process disguise'),
        ('Covert Exfiltration', 'HTTPS-encrypted data transmission disguised as legitimate Discord traffic'),
        ('Automated Persistence', 'Platform-specific autostart mechanisms (Windows Registry, systemd services, Android boot receivers)')
    ]
    
    for feature_name, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature_name}: ').bold = True
        p.add_run(description)
    
    # Educational Value
    doc.add_heading('Educational Value', 1)
    p5 = doc.add_paragraph('This project provides hands-on experience with:')
    
    educational_items = [
        'Low-level system APIs and keyboard hooks (Win32 API, X11/Xorg)',
        'Cryptographic implementations and key management',
        'Computer vision and biometric authentication systems',
        'Network protocols and covert communication channels',
        'Cross-platform software packaging (PyInstaller, Buildozer)',
        'Malware analysis and reverse engineering concepts',
        'Ethical hacking methodologies and responsible disclosure'
    ]
    
    for item in educational_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # Ethical Considerations
    doc.add_heading('Ethical Considerations', 1)
    p6 = doc.add_paragraph()
    p6.add_run('This project was developed strictly for ')
    run = p6.add_run('educational and research purposes')
    run.bold = True
    p6.add_run(' within a controlled environment. All testing was conducted on personally-owned devices with '
               'explicit consent. The project demonstrates the importance of cybersecurity awareness and the need '
               'for robust defensive measures against keystroke logging attacks.')
    
    # Report Structure
    doc.add_heading('Report Structure', 1)
    p7 = doc.add_paragraph(
        'This comprehensive report details the technical architecture, implementation logic, security features, '
        'platform-specific adaptations, attack vectors, defensive countermeasures, and ethical considerations '
        'surrounding this educational keylogger project. The analysis provides valuable insights for both '
        'offensive security research and defensive cybersecurity practices.'
    )
    
    # Save the document
    doc.save('Project_Introduction.docx')
    print("✓ Word document created successfully: Project_Introduction.docx")
    
except ImportError:
    print("ERROR: python-docx is not installed. Installing now...")
    import subprocess
    subprocess.run(['pip', 'install', 'python-docx'])
    print("Please run this script again.")
except Exception as e:
    print(f"ERROR: {e}")
